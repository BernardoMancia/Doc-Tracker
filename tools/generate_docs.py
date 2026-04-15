import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import fitz

BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "documentacao" / "geral"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

BLUE_DARK = RGBColor(0x0F, 0x17, 0x2A)
BLUE_MID = RGBColor(0x1E, 0x40, 0xAF)
BLUE_LIGHT = RGBColor(0x3B, 0x82, 0xF6)
CYAN = RGBColor(0x06, 0xD6, 0xA0)
TEXT_DARK = RGBColor(0x1F, 0x29, 0x37)
TEXT_MED = RGBColor(0x47, 0x55, 0x69)
TEXT_LIGHT = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "0F172A"
TABLE_ALT_BG = "F1F5F9"
TABLE_BORDER = "CBD5E1"


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{val}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = BLUE_MID
    if level == 1:
        run.font.size = Pt(18)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(8)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        run2 = p2.add_run()
        run2.add_break()
    elif level == 2:
        run.font.size = Pt(14)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or TEXT_DARK
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    shade = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
    p._p.get_or_add_pPr().append(shade)
    return p


def add_pro_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, TABLE_HEADER_BG)
        set_cell_border(cell, top=TABLE_BORDER, bottom=TABLE_BORDER, left=TABLE_BORDER, right=TABLE_BORDER)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT_DARK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                shade_cell(cell, TABLE_ALT_BG)
            set_cell_border(cell, top=TABLE_BORDER, bottom=TABLE_BORDER, left=TABLE_BORDER, right=TABLE_BORDER)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table


def add_cover(doc, title_main, subtitle, version_text, footer_text):
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    p_icon = doc.add_paragraph()
    p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_icon = p_icon.add_run("_______________________________________________")
    run_icon.font.color.rgb = CYAN
    run_icon.font.size = Pt(14)
    run_icon.bold = True

    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_brand.paragraph_format.space_before = Pt(20)
    run_b = p_brand.add_run("OSINT & DLP")
    run_b.font.size = Pt(36)
    run_b.bold = True
    run_b.font.color.rgb = BLUE_MID

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sys.paragraph_format.space_before = Pt(4)
    run_s = p_sys.add_run("SYSTEM")
    run_s.font.size = Pt(20)
    run_s.font.color.rgb = CYAN
    run_s.font.name = "Calibri Light"

    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line2.paragraph_format.space_before = Pt(8)
    run_l2 = p_line2.add_run("_______________________________________________")
    run_l2.font.color.rgb = CYAN
    run_l2.font.size = Pt(14)
    run_l2.bold = True

    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = TEXT_MED

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ver.paragraph_format.space_before = Pt(12)
    run_ver = p_ver.add_run(version_text)
    run_ver.font.size = Pt(11)
    run_ver.font.color.rgb = TEXT_LIGHT

    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ft = p_foot.add_run(footer_text)
    run_ft.font.size = Pt(9)
    run_ft.font.color.rgb = TEXT_LIGHT
    run_ft.italic = True

    doc.add_page_break()


def add_toc_page(doc, sections):
    add_heading_styled(doc, sections["toc_title"], 1)
    for i, (num, title) in enumerate(sections["items"]):
        p = doc.add_paragraph()
        run = p.add_run(f"{num}   {title}")
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK if i % 2 == 0 else TEXT_MED
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
    doc.add_page_break()


def build_pt():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_DARK

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover(
        doc,
        "OSINT & DLP SYSTEM",
        "Grupo Roullier \u2014 Documenta\u00e7\u00e3o T\u00e9cnica Completa",
        "Vers\u00e3o 2.1  |  Confidencial \u2014 Uso Interno",
        "Grupo Roullier Security Team  \u00b7  2026",
    )

    add_toc_page(doc, {
        "toc_title": "Sum\u00e1rio",
        "items": [
            ("01", "O Que \u00e9 o OSINT & DLP System"),
            ("02", "Gloss\u00e1rio de Termos"),
            ("03", "Arquitetura e Componentes"),
            ("04", "Fluxo de Opera\u00e7\u00e3o Detalhado"),
            ("05", "Intelligence Matrix"),
            ("06", "Detec\u00e7\u00e3o, Classifica\u00e7\u00e3o e Categoriza\u00e7\u00e3o"),
            ("07", "Modos de Opera\u00e7\u00e3o"),
            ("08", "Instala\u00e7\u00e3o Passo a Passo"),
            ("09", "Configura\u00e7\u00e3o de Vari\u00e1veis de Ambiente"),
            ("10", "Execu\u00e7\u00e3o e Comandos"),
            ("11", "Dashboard Web \u2014 Guia Visual"),
            ("12", "Alertas Telegram"),
            ("13", "API REST \u2014 Refer\u00eancia Completa"),
            ("14", "Seguran\u00e7a e Privacidade"),
            ("15", "Manuten\u00e7\u00e3o e Ingest\u00e3o de Dados"),
        ],
    })

    add_heading_styled(doc, "01. O Que \u00e9 o OSINT & DLP System", 1)
    add_para(doc,
        'O nome do sistema \u00e9 composto por dois acr\u00f4nimos amplamente reconhecidos na \u00e1rea de seguran\u00e7a da informa\u00e7\u00e3o:'
    )
    add_bullet(doc, 'OSINT (Open Source Intelligence) \u2014 Intelig\u00eancia em Fontes Abertas: a pr\u00e1tica de coletar e analisar informa\u00e7\u00f5es dispon\u00edveis publicamente na internet.')
    add_bullet(doc, 'DLP (Data Leak Prevention) \u2014 Preven\u00e7\u00e3o de Vazamento de Dados: conjunto de t\u00e9cnicas para detectar e impedir a exposi\u00e7\u00e3o n\u00e3o autorizada de informa\u00e7\u00f5es sens\u00edveis.')
    add_para(doc,
        'Combinados, o OSINT & DLP System \u00e9 uma plataforma automatizada que varre continuamente a internet p\u00fablica '
        'em busca de documentos, planilhas e arquivos que possam conter informa\u00e7\u00f5es confidenciais do Grupo Roullier. '
        'O sistema monitora as tr\u00eas empresas principais do grupo no Brasil:'
    )
    add_bullet(doc, 'Timac Agro Brasil \u2014 Ind\u00fastria e Com\u00e9rcio de Fertilizantes')
    add_bullet(doc, 'Sulfabr\u00e1s Sulfatos do Brasil \u2014 Fabrica\u00e7\u00e3o de Sulfatos e Qu\u00edmicos')
    add_bullet(doc, 'Phosphea Brasil \u2014 Com\u00e9rcio de Fosfatos')
    add_para(doc,
        'Quando o sistema encontra um documento suspeito, ele o baixa diretamente na mem\u00f3ria do servidor (sem salvar em disco), '
        'extrai o texto, analisa com express\u00f5es regulares, classifica o n\u00edvel de risco e notifica a equipe de seguran\u00e7a. '
        'Tudo isso de forma totalmente autom\u00e1tica, rodando 24/7.'
    )

    add_heading_styled(doc, "02. Gloss\u00e1rio de Termos", 1)
    add_pro_table(doc,
        ["Termo", "Significado"],
        [
            ["Dork", "Consulta de busca avan\u00e7ada (Google Dork) que utiliza operadores especiais para encontrar resultados espec\u00edficos"],
            ["Finding", "Um resultado de varredura \u2014 um documento ou p\u00e1gina web que cont\u00e9m dados sens\u00edveis detectados"],
            ["Scan", "Uma execu\u00e7\u00e3o completa de varredura: gera\u00e7\u00e3o de dorks \u2192 busca \u2192 download \u2192 inspe\u00e7\u00e3o \u2192 classifica\u00e7\u00e3o"],
            ["Score", "Pontua\u00e7\u00e3o num\u00e9rica (1-100) atribu\u00edda a cada finding indicando a gravidade do risco"],
            ["Snippet", "Trecho mascarado do texto onde dados sens\u00edveis foram encontrados"],
            ["Triagem", "Processo de analisar e classificar cada finding como leg\u00edtimo, falso positivo ou resolvido"],
            ["Intelligence Matrix", "Base de conhecimento com todas as entidades, CNPJs, CPFs e termos monitorados"],
            ["Crawler", "Motor de busca autom\u00e1tica que executa as consultas e coleta URLs"],
            ["Inspector", "M\u00f3dulo que baixa, extrai texto, inspeciona e classifica documentos"],
            ["Soft Delete", "M\u00e9todo de exclus\u00e3o que marca registros como inativos sem remov\u00ea-los fisicamente"],
            ["SSE", "Server-Sent Events \u2014 tecnologia para enviar notifica\u00e7\u00f5es em tempo real do servidor para o navegador"],
            ["TLD", "Top-Level Domain \u2014 sufixo do dom\u00ednio (.com.br, .pt, .fr) usado para detectar pa\u00eds de origem"],
        ],
        col_widths=[4, 13],
    )

    add_heading_styled(doc, "03. Arquitetura e Componentes", 1)
    add_para(doc,
        'O sistema \u00e9 constru\u00eddo em Python 3.10+ e segue uma arquitetura modular em camadas independentes. '
        'Dois modos de opera\u00e7\u00e3o (Server e Desktop) compartilham o mesmo n\u00facleo de l\u00f3gica, garantindo resultados id\u00eanticos.'
    )
    add_heading_styled(doc, "3.1 Mapa de Componentes", 2)
    add_pro_table(doc,
        ["M\u00f3dulo", "Pasta", "Fun\u00e7\u00e3o", "Tecnologia"],
        [
            ["Configura\u00e7\u00e3o", "config/", "Settings globais e Intelligence Matrix", "Pydantic Settings"],
            ["Core", "core/", "Modelos de dados e engine de banco", "SQLAlchemy Async + aiosqlite"],
            ["Crawler", "crawler/", "Gera\u00e7\u00e3o de dorks e motor de busca", "DuckDuckGo Search (ddgs)"],
            ["Inspector", "inspector/", "Download, extra\u00e7\u00e3o de texto, regex, risco", "httpx, PyMuPDF, python-docx, openpyxl"],
            ["Alertas", "alerts/", "Notifica\u00e7\u00f5es Telegram", "httpx"],
            ["API", "api/", "Dashboard web e API REST", "FastAPI + Jinja2 + Chart.js"],
            ["Server", "server/", "Entry point do modo servidor", "Uvicorn + APScheduler"],
            ["Desktop", "desktop/", "Interface gr\u00e1fica standalone", "Tkinter"],
            ["Tools", "tools/", "Scripts utilit\u00e1rios (ingest\u00e3o, documenta\u00e7\u00e3o)", "Python stdlib"],
            ["Data", "data/", "Entidades extra\u00eddas dos cadastros", "JSON"],
        ],
        col_widths=[3, 2.5, 6.5, 5],
    )

    add_heading_styled(doc, "3.2 Estrutura de Pastas", 2)
    add_code(doc,
        "Doc-Tracker/\n"
        "+-- config/                  Configuracoes compartilhadas\n"
        "|   +-- settings.py          Variaveis de ambiente\n"
        "|   +-- intelligence_matrix.py  Matriz de inteligencia\n"
        "+-- core/                    Banco de dados\n"
        "|   +-- database.py          Engine async SQLite\n"
        "|   +-- models.py            Modelos Scan e Finding\n"
        "+-- crawler/                 Motor de busca\n"
        "|   +-- dork_generator.py    80+ consultas automaticas\n"
        "|   +-- search_engine.py     Busca via DuckDuckGo\n"
        "|   +-- url_filter.py        Filtro de dominios e deteccao\n"
        "+-- inspector/               Inspecao de documentos\n"
        "|   +-- downloader.py        Download em memoria\n"
        "|   +-- extractor.py         Extracao PDF/DOCX/XLSX/TXT\n"
        "|   +-- regex_engine.py      Deteccao de PII por regex\n"
        "|   +-- risk_classifier.py   Classificacao de risco\n"
        "+-- alerts/                  Webhooks\n"
        "|   +-- webhook.py           Telegram alerts\n"
        "+-- api/                     Dashboard e API\n"
        "|   +-- app.py               FastAPI application\n"
        "|   +-- routes/              Endpoints REST\n"
        "|   +-- templates/           HTML do dashboard\n"
        "+-- static/                  Frontend\n"
        "|   +-- css/dashboard.css    Estilos dark mode\n"
        "|   +-- js/dashboard.js      Logica do dashboard\n"
        "+-- server/                  Modo servidor\n"
        "|   +-- run_server.py        Entry point + scheduler\n"
        "|   +-- osint-dlp.service    Servico systemd\n"
        "+-- desktop/                 Modo desktop\n"
        "|   +-- app.py               GUI Tkinter\n"
        "|   +-- run_desktop.py       Entry point\n"
        "+-- data/entities.json       Dados das entidades\n"
        "+-- documentacao/            Esta documentacao\n"
        "+-- requirements.txt         Dependencias Python\n"
        "+-- .env.example             Template de configuracao\n"
        "+-- run.py                   Entry point dev local"
    )

    add_heading_styled(doc, "04. Fluxo de Opera\u00e7\u00e3o Detalhado", 1)
    add_para(doc, 'Cada scan completo segue esta sequ\u00eancia de 7 etapas:')
    steps = [
        ("ETAPA 1 \u2014 Gera\u00e7\u00e3o de Dorks",
         "O DorkGenerator cria aproximadamente 80 consultas de busca avan\u00e7adas. Cada dork combina "
         "nomes de empresas, CNPJs, CPFs, fornecedores e termos sens\u00edveis com operadores de busca "
         "para plataformas espec\u00edficas (Scribd, SlideShare, Issuu, GitHub, GitLab, Google Drive). "
         "Dom\u00ednios oficiais do grupo s\u00e3o automaticamente exclu\u00eddos com o operador -site:."),
        ("ETAPA 2 \u2014 Varredura Web",
         "O SearchEngine executa cada dork via DuckDuckGo com delays randomizados de 2 a 4 segundos "
         "entre consultas, evitando bloqueio por rate-limiting. Cada resultado retorna t\u00edtulo, URL e snippet."),
        ("ETAPA 3 \u2014 Deduplica\u00e7\u00e3o",
         "Antes de processar cada URL, o sistema verifica no banco de dados se ela j\u00e1 foi analisada "
         "anteriormente. URLs j\u00e1 conhecidas s\u00e3o ignoradas, garantindo que apenas novos documentos sejam processados."),
        ("ETAPA 4 \u2014 Download em Mem\u00f3ria",
         "O Downloader obt\u00e9m cada documento diretamente em mem\u00f3ria (BytesIO) via httpx, sem gravar "
         "nenhum arquivo no disco. Isso minimiza a pegada forense e acelera o processamento."),
        ("ETAPA 5 \u2014 Extra\u00e7\u00e3o de Texto e Metadados",
         "O TextExtractor identifica o tipo de arquivo e processa: PDF via PyMuPDF, DOCX via python-docx, "
         "XLSX via openpyxl, e texto puro. Al\u00e9m do conte\u00fado, extrai metadados como autor, software criador "
         "e data de cria\u00e7\u00e3o, filtrando artefatos t\u00e9cnicos (nomes de ferramentas como Microsoft Office)."),
        ("ETAPA 6 \u2014 Inspe\u00e7\u00e3o por Regex e Classifica\u00e7\u00e3o",
         "O RegexEngine analisa o texto em busca de CPFs, CNPJs, dados financeiros, entidades do grupo "
         "e termos sens\u00edveis. Os snippets s\u00e3o mascarados para proteger PII. O RiskClassifier ent\u00e3o "
         "atribui um score num\u00e9rico (1-100) e n\u00edvel (Cr\u00edtico/Alto/M\u00e9dio/Baixo)."),
        ("ETAPA 7 \u2014 Persist\u00eancia e Alertas",
         "Os findings s\u00e3o salvos no banco SQLite. O pa\u00eds de origem \u00e9 detectado pelo TLD do dom\u00ednio e "
         "o idioma por an\u00e1lise textual. Ao final do scan, um resumo com todos os novos links \u00e9 enviado "
         "via Telegram, agrupado por n\u00edvel de risco."),
    ]
    for title, desc in steps:
        add_heading_styled(doc, title, 3)
        add_para(doc, desc)

    add_heading_styled(doc, "05. Intelligence Matrix", 1)
    add_para(doc,
        'A Intelligence Matrix \u00e9 a base de conhecimento central do sistema. Ela \u00e9 carregada dinamicamente '
        'do arquivo data/entities.json, que cont\u00e9m dados extra\u00eddos das fichas cadastrais do grupo.'
    )
    add_para(doc, 'Dados monitorados na vers\u00e3o atual:', bold=True)
    add_pro_table(doc,
        ["Categoria", "Quantidade", "Exemplos"],
        [
            ["Entidades (Raz\u00f5es Sociais)", "19 nomes", "TIMAC AGRO IND\u00daTRIA E COM\u00c9RCIO DE FERTILIZANTES LTDA"],
            ["CNPJs", "24 n\u00fameros", "02.329.713/0001-29, 26.769.908/0001-58"],
            ["CPFs", "Administradores", "Soci\u00f3s e diretores do grupo"],
            ["Fornecedores", "21 empresas", "OCP Brasil, Pisani Pl\u00e1sticos, MOVIDA"],
            ["Pessoas-chave", "4 nomes", "Administradores e s\u00f3cios registrados"],
            ["Emails Corporativos", "9 endere\u00e7os", "nfe1@timacagro.com.br"],
            ["Termos Sens\u00edveis", "15+ termos", "confidencial, uso interno, senha, password"],
        ],
        col_widths=[5, 3, 9],
    )
    add_para(doc,
        'Para atualizar a matrix com novas entidades, basta adicionar novos PDFs cadastrais na pasta '
        'anexos/ e executar o script de ingest\u00e3o (Se\u00e7\u00e3o 15).'
    )

    add_heading_styled(doc, "06. Detec\u00e7\u00e3o, Classifica\u00e7\u00e3o e Categoriza\u00e7\u00e3o", 1)
    add_heading_styled(doc, "6.1 N\u00edveis de Risco", 2)
    add_pro_table(doc,
        ["N\u00edvel", "Emoji", "Score", "Crit\u00e9rios de Classifica\u00e7\u00e3o"],
        [
            ["Cr\u00edtico", "vermelho", "80-100", "M\u00faltiplos CPFs + CNPJs + dados financeiros, ou CPFs expostos com contexto altamente sens\u00edvel"],
            ["Alto", "laranja", "60-79", "CPFs ou CNPJs do grupo combinados com termos financeiros"],
            ["M\u00e9dio", "amarelo", "40-59", "Men\u00e7\u00e3o a entidades do grupo com termos sens\u00edveis moderados"],
            ["Baixo", "verde", "1-39", "Men\u00e7\u00e3o gen\u00e9rica a entidades ou termos isolados sem contexto cr\u00edtico"],
        ],
        col_widths=[2, 2, 2, 11],
    )
    add_heading_styled(doc, "6.2 Categorias de Finding", 2)
    add_pro_table(doc,
        ["Categoria", "Descri\u00e7\u00e3o", "Exemplo de Detec\u00e7\u00e3o"],
        [
            ["RH", "Dados de funcion\u00e1rios e trabalhistas", "Folha de pagamento com CPFs de colaboradores"],
            ["Financeiro", "Dados monet\u00e1rios e cont\u00e1beis", "Balan\u00e7o patrimonial, faturamento, contratos"],
            ["TI", "Infraestrutura tecnol\u00f3gica", "Configura\u00e7\u00f5es de rede, diagramas de sistema"],
            ["TI/Seguran\u00e7a", "Credenciais e acessos", "Senhas, tokens de API, chaves VPN"],
            ["Dados Pessoais", "PII sem contexto corporativo", "CPFs e endere\u00e7os em listagens p\u00fablicas"],
            ["Corporativo", "Documentos estrat\u00e9gicos", "Propostas comerciais, contratos de fornecimento"],
        ],
        col_widths=[3, 5, 9],
    )
    add_heading_styled(doc, "6.3 Filtros Regionais", 2)
    add_para(doc,
        'O sistema detecta automaticamente o pa\u00eds de origem pela extens\u00e3o do dom\u00ednio (TLD) e o idioma '
        'do conte\u00fado por an\u00e1lise textual. Isso permite filtrar findings por regi\u00e3o no dashboard.'
    )
    add_pro_table(doc,
        ["Pa\u00eds", "TLDs Detectados", "Idiomas"],
        [
            ["Brasil", ".com.br, .org.br, .gov.br", "Portugu\u00eas (pt)"],
            ["Portugal", ".pt", "Portugu\u00eas (pt)"],
            ["Fran\u00e7a", ".fr", "Franc\u00eas (fr)"],
            ["EUA", ".us, .gov", "Ingl\u00eas (en)"],
            ["Espanha", ".es", "Espanhol (es)"],
            ["Internacional", "Todos os demais", "Detectado por an\u00e1lise"],
        ],
        col_widths=[3, 5.5, 8.5],
    )

    add_heading_styled(doc, "07. Modos de Opera\u00e7\u00e3o", 1)
    add_heading_styled(doc, "7.1 Modo Server (Produ\u00e7\u00e3o)", 2)
    add_para(doc,
        'O modo Server \u00e9 a vers\u00e3o principal, projetada para ambientes de produ\u00e7\u00e3o com monitoramento 24/7.'
    )
    add_para(doc, 'Recursos exclusivos:', bold=True)
    for item in [
        'Dashboard web com tema cybersecurity dark mode (glassmorphism, neon, gradientes)',
        'Gr\u00e1ficos interativos: distribui\u00e7\u00e3o por risco (doughnut) e fontes de vazamento (barras)',
        'Tabela de triagem com filtros por risco, status, categoria, pa\u00eds e idioma',
        'Barra de progresso em tempo real com percentual durante scans',
        'Agendamento autom\u00e1tico via APScheduler (padr\u00e3o: 08h, 14h, 22h BRT)',
        'Alertas Telegram: startup do sistema, in\u00edcio de scan e resumo com links novos',
        'SSE (Server-Sent Events) para push notifications no dashboard',
        'API REST completa para integra\u00e7\u00f5es externas',
        'Servi\u00e7o systemd para opera\u00e7\u00e3o cont\u00ednua sem login (Linux)',
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "7.2 Modo Desktop (Consulta Ad-hoc)", 2)
    add_para(doc,
        'O modo Desktop \u00e9 uma aplica\u00e7\u00e3o GUI standalone constru\u00edda com Tkinter, para consultas pontuais '
        'sem necessidade de servidor web.'
    )
    for item in [
        'Interface gr\u00e1fica dark mode com barra de progresso e m\u00e9tricas',
        'Exporta\u00e7\u00e3o de resultados para CSV e JSON',
        'Opera\u00e7\u00e3o independente \u2014 n\u00e3o precisa de servidor web rodando',
        'Compil\u00e1vel para .exe standalone via cx_Freeze',
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "08. Instala\u00e7\u00e3o Passo a Passo", 1)
    add_heading_styled(doc, "8.1 Requisitos do Sistema", 2)
    add_pro_table(doc,
        ["Requisito", "M\u00ednimo", "Recomendado"],
        [
            ["Python", "3.10", "3.12+"],
            ["Sistema Operacional", "Linux ou Windows", "Ubuntu 22.04 LTS (Server)"],
            ["RAM", "512 MB", "1 GB"],
            ["Disco", "100 MB", "500 MB (banco pode crescer)"],
            ["Rede", "Acesso \u00e0 internet", "Sem restri\u00e7\u00f5es de firewall para HTTPS"],
        ],
        col_widths=[4, 5, 8],
    )

    add_heading_styled(doc, "8.2 Passo 1: Clonar o Reposit\u00f3rio", 2)
    add_code(doc, "git clone <url-do-repositorio>\ncd Doc-Tracker")

    add_heading_styled(doc, "8.3 Passo 2: Criar Ambiente Virtual", 2)
    add_para(doc, 'Linux / macOS:')
    add_code(doc, "python3 -m venv venv\nsource venv/bin/activate")
    add_para(doc, 'Windows:')
    add_code(doc, "python -m venv venv\nvenv\\Scripts\\activate")

    add_heading_styled(doc, "8.4 Passo 3: Instalar Depend\u00eancias", 2)
    add_code(doc, "pip install -r requirements.txt")

    add_heading_styled(doc, "8.5 Passo 4: Configurar Vari\u00e1veis de Ambiente", 2)
    add_code(doc, "cp .env.example .env\nnano .env   # editar com valores reais")
    add_para(doc, 'Veja a Se\u00e7\u00e3o 09 para detalhes de cada vari\u00e1vel.')

    add_heading_styled(doc, "8.6 Passo 5: Executar", 2)
    add_para(doc, 'Escolha o modo de execu\u00e7\u00e3o adequado (Se\u00e7\u00e3o 10).')

    add_heading_styled(doc, "8.7 Passo 6 (Opcional): Instalar como Servi\u00e7o", 2)
    add_para(doc, 'Para servidores Linux em produ\u00e7\u00e3o, instale como servi\u00e7o systemd:')
    add_code(doc,
        "# Editar o arquivo de servico para ajustar caminhos e usuario:\n"
        "nano server/osint-dlp.service\n\n"
        "# Copiar e ativar:\n"
        "sudo cp server/osint-dlp.service /etc/systemd/system/\n"
        "sudo systemctl daemon-reload\n"
        "sudo systemctl enable osint-dlp\n"
        "sudo systemctl start osint-dlp"
    )
    add_para(doc, 'Verificar status:')
    add_code(doc, "sudo systemctl status osint-dlp\nsudo journalctl -u osint-dlp -f")

    add_heading_styled(doc, "09. Configura\u00e7\u00e3o de Vari\u00e1veis de Ambiente", 1)
    add_para(doc,
        'O arquivo .env cont\u00e9m todas as configura\u00e7\u00f5es sens\u00edveis. Ele NUNCA deve ser enviado ao reposit\u00f3rio Git '
        '(j\u00e1 est\u00e1 no .gitignore). Use o .env.example como template.'
    )
    add_pro_table(doc,
        ["Vari\u00e1vel", "Descri\u00e7\u00e3o", "Padr\u00e3o", "Obrigat\u00f3ria"],
        [
            ["DATABASE_URL", "URL de conex\u00e3o com o banco SQLite", "sqlite+aiosqlite:///osint_dlp.db", "Sim"],
            ["API_HOST", "Host em que o servidor escuta", "0.0.0.0", "N\u00e3o"],
            ["API_PORT", "Porta do servidor web", "8443", "N\u00e3o"],
            ["TELEGRAM_BOT_TOKEN", "Token do bot Telegram (obtido via @BotFather)", "\u2014", "Para alertas"],
            ["TELEGRAM_CHAT_ID", "ID do chat/grupo Telegram para receber alertas", "\u2014", "Para alertas"],
            ["SCAN_SCHEDULE_HOURS", "Hor\u00e1rios de scan autom\u00e1tico (separados por v\u00edrgula)", "8,14,22", "N\u00e3o"],
            ["SERVER_MODE", "Ativar funcionalidades de servidor (scheduler, alerts)", "true", "N\u00e3o"],
            ["DASHBOARD_URL", "URL p\u00fablica do dashboard (usada nos links do Telegram)", "http://localhost:8443", "N\u00e3o"],
            ["PROXY_URL", "Proxy SOCKS5 para crawling (ex: socks5://host:port)", "\u2014", "N\u00e3o"],
        ],
        col_widths=[4, 5.5, 4.5, 3],
    )

    add_heading_styled(doc, "10. Execu\u00e7\u00e3o e Comandos", 1)
    add_pro_table(doc,
        ["Modo", "Comando", "Descri\u00e7\u00e3o"],
        [
            ["Desenvolvimento Local", "python run.py", "Servidor web local sem scheduler nem Telegram"],
            ["Servidor (Produ\u00e7\u00e3o)", "python server/run_server.py", "Dashboard + Telegram + Auto-scan 3x/dia"],
            ["Desktop (GUI)", "python desktop/run_desktop.py", "Interface gr\u00e1fica standalone"],
            ["Compilar Desktop", "python setup_desktop.py build", "Gera execut\u00e1vel .exe via cx_Freeze"],
            ["Ingest\u00e3o de Dados", "python tools/ingest_anexos.py", "Processa novos PDFs cadastrais"],
            ["Gerar Documenta\u00e7\u00e3o", "python tools/generate_docs.py", "Recria estes documentos"],
        ],
        col_widths=[4, 6, 7],
    )

    add_heading_styled(doc, "11. Dashboard Web \u2014 Guia Visual", 1)
    add_para(doc,
        'O dashboard segue padr\u00f5es de design de ferramentas de cybersecurity profissionais: '
        'tema dark mode, glassmorphism, cores neon para status, tipografia Inter e JetBrains Mono.'
    )
    add_heading_styled(doc, "11.1 Se\u00e7\u00f5es do Dashboard", 2)
    add_pro_table(doc,
        ["Se\u00e7\u00e3o", "Descri\u00e7\u00e3o"],
        [
            ["Header", "Logo, status do sistema (Idle/Scanning com dot pulsante), bot\u00e3o 'Iniciar Scan'"],
            ["Barra de Progresso", "Aparece durante scans com fase atual, percentual e detalhe do dork/URL"],
            ["M\u00e9tricas", "5 cards com Total Findings, Cr\u00edtico, Alto, M\u00e9dio e Baixo (hover com anima\u00e7\u00e3o)"],
            ["Gr\u00e1fico de Risco", "Doughnut chart mostrando distribui\u00e7\u00e3o por n\u00edvel de risco"],
            ["Fontes de Vazamento", "Bar chart horizontal com as plataformas que mais cont\u00eam vazamentos"],
            ["Tabela de Triagem", "Lista completa de findings com 13 colunas filtr\u00e1veis"],
            ["Painel de Detalhes", "Slide-in lateral com informa\u00e7\u00f5es completas do finding selecionado"],
            ["A\u00e7\u00f5es", "Bot\u00f5es para classificar: Investigar, Falso Positivo, Resolvido, Notificado"],
            ["Pagina\u00e7\u00e3o", "Navega\u00e7\u00e3o entre p\u00e1ginas de resultados (20 por p\u00e1gina)"],
            ["Toasts", "Notifica\u00e7\u00f5es flutuantes para confirma\u00e7\u00f5es de a\u00e7\u00f5es e novos alertas"],
        ],
        col_widths=[4, 13],
    )

    add_heading_styled(doc, "12. Alertas Telegram", 1)
    add_para(doc, 'O sistema envia 3 tipos de alertas via Telegram:')
    add_heading_styled(doc, "12.1 Tipos de Alerta", 2)
    add_pro_table(doc,
        ["Alerta", "Quando", "Conte\u00fado"],
        [
            ["Sistema Online", "Ao iniciar o servi\u00e7o", "URL do dashboard, hor\u00e1rio, schedule de scans configurado"],
            ["Scan Iniciado", "In\u00edcio de cada scan", "N\u00famero de dorks a executar, timestamp"],
            ["Resumo do Scan", "Final de cada scan", "Contagem por n\u00edvel + lista de TODOS os novos links agrupados por risco"],
        ],
        col_widths=[3.5, 4, 9.5],
    )
    add_heading_styled(doc, "12.2 Comportamento Incremental", 2)
    add_para(doc,
        'O sistema utiliza deduplica\u00e7\u00e3o por URL: se um documento j\u00e1 foi encontrado em um scan anterior, '
        'ele n\u00e3o \u00e9 reprocessado nem notificado novamente. Isso garante que apenas NOVOS vazamentos '
        'geram alertas nos scans subsequentes.'
    )

    add_heading_styled(doc, "13. API REST \u2014 Refer\u00eancia Completa", 1)
    add_pro_table(doc,
        ["M\u00e9todo", "Endpoint", "Descri\u00e7\u00e3o"],
        [
            ["GET", "/", "Dashboard HTML completo"],
            ["GET", "/api/dashboard", "M\u00e9tricas agregadas: totais, por risco, por plataforma, por categoria, \u00faltimo scan"],
            ["GET", "/api/findings?page=N&risk_level=X&status=Y&category=Z&country=C&language=L", "Listar findings com pagina\u00e7\u00e3o e filtros"],
            ["PATCH", "/api/findings/{id}/status", "Atualizar status: pending, investigating, false_positive, resolved, notified"],
            ["DELETE", "/api/findings/{id}", "Soft delete (marca is_deleted=True, mant\u00e9m no banco)"],
            ["GET", "/api/scans", "Hist\u00f3rico dos \u00faltimos 50 scans"],
            ["POST", "/api/scans/trigger", "Disparar scan manual (rejeita se j\u00e1 houver um rodando)"],
            ["GET", "/api/scans/progress", "Progresso atual: fase, current, total, detalhe"],
            ["GET", "/api/stream", "Stream SSE com eventos new_finding em tempo real"],
        ],
        col_widths=[1.5, 8, 7.5],
    )

    add_heading_styled(doc, "14. Seguran\u00e7a e Privacidade", 1)
    add_pro_table(doc,
        ["Medida", "Descri\u00e7\u00e3o"],
        [
            ["Download em Mem\u00f3ria", "Documentos s\u00e3o baixados em BytesIO e processados sem tocar o filesystem"],
            ["Mascaramento de PII", "CPFs, CNPJs e emails s\u00e3o mascarados nos snippets (ex: ***.456.789-**)"],
            ["Dom\u00ednios Exclu\u00eddos", "Sites oficiais (timacagro.com.br, phosphea.com, roullier.com) nunca s\u00e3o varridos"],
            ["Soft Delete", "Dados nunca s\u00e3o removidos fisicamente \u2014 marcados como inativos para auditoria"],
            [".env Protegido", "Credenciais sens\u00edveis ficam fora do reposit\u00f3rio via .gitignore"],
            ["Delays Anti-bloqueio", "Intervalos randomizados (2-4s) entre requisi\u00e7\u00f5es de busca"],
            ["Dedup por URL", "Cada URL \u00e9 processada apenas uma vez em todo o hist\u00f3rico"],
        ],
        col_widths=[4, 13],
    )

    add_heading_styled(doc, "15. Manuten\u00e7\u00e3o e Ingest\u00e3o de Dados", 1)
    add_para(doc, 'Para adicionar novas empresas ou atualizar dados cadastrais:', bold=True)
    add_para(doc, '1. Coloque os PDFs das fichas cadastrais na pasta anexos/ (na raiz do projeto).')
    add_para(doc, '2. Execute o script de ingest\u00e3o:')
    add_code(doc, "python tools/ingest_anexos.py")
    add_para(doc, '3. O script extrai automaticamente:')
    for item in [
        'Raz\u00e3o social e CNPJ principal',
        'Inscri\u00e7\u00e3o estadual, data de funda\u00e7\u00e3o, capital social',
        'S\u00f3cios com percentual de participa\u00e7\u00e3o',
        'Administradores com CPFs',
        'Filiais com CNPJs e endere\u00e7os',
        'Todos os CNPJs e CPFs mencionados no documento',
        'Emails, telefones, fornecedores e refer\u00eancias banc\u00e1rias',
    ]:
        add_bullet(doc, item)
    add_para(doc, '4. Os dados s\u00e3o salvos em data/entities.json e os PDFs originais s\u00e3o deletados.')
    add_para(doc, '5. O pr\u00f3ximo scan usar\u00e1 automaticamente as novas entidades.')

    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_e = p_end.add_run("_______________________________________________")
    run_e.font.color.rgb = CYAN
    run_e.bold = True
    p_end2 = doc.add_paragraph()
    p_end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end2.paragraph_format.space_before = Pt(16)
    run_e2 = p_end2.add_run("Documento Confidencial\nGrupo Roullier Security Team\nVers\u00e3o 2.1 \u2014 2026")
    run_e2.font.size = Pt(10)
    run_e2.font.color.rgb = TEXT_LIGHT
    run_e2.italic = True

    return doc


def build_en():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_DARK

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover(
        doc,
        "OSINT & DLP SYSTEM",
        "Roullier Group \u2014 Complete Technical Documentation",
        "Version 2.1  |  Confidential \u2014 Internal Use Only",
        "Roullier Group Security Team  \u00b7  2026",
    )

    add_toc_page(doc, {
        "toc_title": "Table of Contents",
        "items": [
            ("01", "What is the OSINT & DLP System"),
            ("02", "Glossary of Terms"),
            ("03", "Architecture and Components"),
            ("04", "Detailed Operation Flow"),
            ("05", "Intelligence Matrix"),
            ("06", "Detection, Classification and Categorization"),
            ("07", "Operation Modes"),
            ("08", "Step-by-Step Installation"),
            ("09", "Environment Variables Configuration"),
            ("10", "Execution and Commands"),
            ("11", "Web Dashboard \u2014 Visual Guide"),
            ("12", "Telegram Alerts"),
            ("13", "REST API \u2014 Complete Reference"),
            ("14", "Security and Privacy"),
            ("15", "Maintenance and Data Ingestion"),
        ],
    })

    add_heading_styled(doc, "01. What is the OSINT & DLP System", 1)
    add_para(doc,
        'The system name is composed of two widely recognized acronyms in information security:'
    )
    add_bullet(doc, 'OSINT (Open Source Intelligence) \u2014 The practice of collecting and analyzing publicly available information from the internet.')
    add_bullet(doc, 'DLP (Data Leak Prevention) \u2014 A set of techniques to detect and prevent unauthorized exposure of sensitive information.')
    add_para(doc,
        'Combined, the OSINT & DLP System is an automated platform that continuously scans the public internet '
        'for documents, spreadsheets, and files that may contain confidential information from the Roullier Group. '
        'The system monitors the three main group companies in Brazil:'
    )
    add_bullet(doc, 'Timac Agro Brasil \u2014 Fertilizer Industry and Commerce')
    add_bullet(doc, 'Sulfabr\u00e1s Sulfatos do Brasil \u2014 Sulfate and Chemical Manufacturing')
    add_bullet(doc, 'Phosphea Brasil \u2014 Phosphate Commerce')
    add_para(doc,
        'When the system finds a suspicious document, it downloads it directly into server memory (no disk writes), '
        'extracts the text, analyzes it with regular expressions, classifies the risk level, and notifies the security team. '
        'All of this fully automated, running 24/7.'
    )

    add_heading_styled(doc, "02. Glossary of Terms", 1)
    add_pro_table(doc,
        ["Term", "Meaning"],
        [
            ["Dork", "Advanced search query (Google Dork) using special operators to find specific results"],
            ["Finding", "A scan result \u2014 a document or web page containing detected sensitive data"],
            ["Scan", "A complete scan execution: dork generation \u2192 search \u2192 download \u2192 inspection \u2192 classification"],
            ["Score", "Numeric score (1-100) assigned to each finding indicating risk severity"],
            ["Snippet", "Masked text excerpt where sensitive data was found"],
            ["Triage", "Process of analyzing and classifying each finding as legitimate, false positive, or resolved"],
            ["Intelligence Matrix", "Knowledge base with all monitored entities, CNPJs, CPFs, and terms"],
            ["Crawler", "Automated search engine that executes queries and collects URLs"],
            ["Inspector", "Module that downloads, extracts text, inspects, and classifies documents"],
            ["Soft Delete", "Deletion method that marks records as inactive without physically removing them"],
            ["SSE", "Server-Sent Events \u2014 technology for sending real-time notifications from server to browser"],
            ["TLD", "Top-Level Domain \u2014 domain suffix (.com.br, .pt, .fr) used to detect country of origin"],
        ],
        col_widths=[4, 13],
    )

    add_heading_styled(doc, "03. Architecture and Components", 1)
    add_para(doc,
        'The system is built with Python 3.10+ and follows a modular architecture in independent layers. '
        'Two operation modes (Server and Desktop) share the same core logic, ensuring identical results.'
    )
    add_heading_styled(doc, "3.1 Component Map", 2)
    add_pro_table(doc,
        ["Module", "Folder", "Function", "Technology"],
        [
            ["Configuration", "config/", "Global settings and Intelligence Matrix", "Pydantic Settings"],
            ["Core", "core/", "Data models and database engine", "SQLAlchemy Async + aiosqlite"],
            ["Crawler", "crawler/", "Dork generation and search engine", "DuckDuckGo Search (ddgs)"],
            ["Inspector", "inspector/", "Download, text extraction, regex, risk", "httpx, PyMuPDF, python-docx, openpyxl"],
            ["Alerts", "alerts/", "Telegram notifications", "httpx"],
            ["API", "api/", "Web dashboard and REST API", "FastAPI + Jinja2 + Chart.js"],
            ["Server", "server/", "Server mode entry point", "Uvicorn + APScheduler"],
            ["Desktop", "desktop/", "Standalone GUI interface", "Tkinter"],
            ["Tools", "tools/", "Utility scripts (ingestion, documentation)", "Python stdlib"],
            ["Data", "data/", "Entities extracted from registrations", "JSON"],
        ],
        col_widths=[3, 2.5, 6.5, 5],
    )

    add_heading_styled(doc, "04. Detailed Operation Flow", 1)
    add_para(doc, 'Each complete scan follows this 7-step sequence:')
    steps_en = [
        ("STEP 1 \u2014 Dork Generation",
         "The DorkGenerator creates approximately 80 advanced search queries. Each dork combines "
         "company names, CNPJs, CPFs, suppliers, and sensitive terms with search operators for specific "
         "platforms (Scribd, SlideShare, Issuu, GitHub, GitLab, Google Drive). Official group domains are "
         "automatically excluded using the -site: operator."),
        ("STEP 2 \u2014 Web Crawling",
         "The SearchEngine executes each dork via DuckDuckGo with randomized delays of 2 to 4 seconds "
         "between queries, avoiding rate-limiting blocks. Each result returns title, URL, and snippet."),
        ("STEP 3 \u2014 Deduplication",
         "Before processing each URL, the system checks the database to see if it was previously analyzed. "
         "Already-known URLs are skipped, ensuring only new documents are processed."),
        ("STEP 4 \u2014 In-Memory Download",
         "The Downloader fetches each document directly into memory (BytesIO) via httpx, without writing "
         "any files to disk. This minimizes forensic footprint and speeds up processing."),
        ("STEP 5 \u2014 Text and Metadata Extraction",
         "The TextExtractor identifies the file type and processes: PDF via PyMuPDF, DOCX via python-docx, "
         "XLSX via openpyxl, and plain text. Beyond content, it extracts metadata such as author, creator "
         "software, and creation date, filtering out technical artifacts."),
        ("STEP 6 \u2014 Regex Inspection and Classification",
         "The RegexEngine analyzes the text for CPFs, CNPJs, financial data, group entities, and sensitive "
         "terms. Snippets are masked to protect PII. The RiskClassifier then assigns a numeric score (1-100) "
         "and level (Critical/High/Medium/Low)."),
        ("STEP 7 \u2014 Persistence and Alerts",
         "Findings are saved to the SQLite database. Country of origin is detected by domain TLD and "
         "language by text analysis. At the end of the scan, a summary with all new links is sent via "
         "Telegram, grouped by risk level."),
    ]
    for title, desc in steps_en:
        add_heading_styled(doc, title, 3)
        add_para(doc, desc)

    add_heading_styled(doc, "05. Intelligence Matrix", 1)
    add_para(doc,
        'The Intelligence Matrix is the central knowledge base of the system. It is dynamically loaded '
        'from the data/entities.json file, containing data extracted from the group\'s registration documents.'
    )
    add_pro_table(doc,
        ["Category", "Quantity", "Examples"],
        [
            ["Entities (Company Names)", "19 names", "TIMAC AGRO IND\u00daTRIA E COM\u00c9RCIO DE FERTILIZANTES LTDA"],
            ["CNPJs", "24 numbers", "02.329.713/0001-29, 26.769.908/0001-58"],
            ["CPFs", "Administrators", "Group partners and directors"],
            ["Suppliers", "21 companies", "OCP Brasil, Pisani Pl\u00e1sticos, MOVIDA"],
            ["Key People", "4 names", "Registered administrators and partners"],
            ["Corporate Emails", "9 addresses", "nfe1@timacagro.com.br"],
            ["Sensitive Terms", "15+ terms", "confidential, internal use, password, senha"],
        ],
        col_widths=[5, 3, 9],
    )

    add_heading_styled(doc, "06. Detection, Classification and Categorization", 1)
    add_heading_styled(doc, "6.1 Risk Levels", 2)
    add_pro_table(doc,
        ["Level", "Indicator", "Score", "Classification Criteria"],
        [
            ["Critical", "Red", "80-100", "Multiple CPFs + CNPJs + financial data, or exposed CPFs in highly sensitive context"],
            ["High", "Orange", "60-79", "Group CPFs or CNPJs combined with financial terms"],
            ["Medium", "Yellow", "40-59", "Group entity mentions with moderate sensitive terms"],
            ["Low", "Green", "1-39", "Generic entity mentions or isolated terms without critical context"],
        ],
        col_widths=[2, 2, 2, 11],
    )
    add_heading_styled(doc, "6.2 Finding Categories", 2)
    add_pro_table(doc,
        ["Category", "Description", "Detection Example"],
        [
            ["HR", "Employee and labor data", "Payroll with employee CPFs"],
            ["Financial", "Monetary and accounting data", "Balance sheets, revenue, contracts"],
            ["IT", "Technology infrastructure", "Network configs, system diagrams"],
            ["IT/Security", "Credentials and access", "Passwords, API tokens, VPN keys"],
            ["Personal Data", "PII without corporate context", "CPFs and addresses in public listings"],
            ["Corporate", "Strategic documents", "Commercial proposals, supply contracts"],
        ],
        col_widths=[3, 5, 9],
    )

    add_heading_styled(doc, "07. Operation Modes", 1)
    add_heading_styled(doc, "7.1 Server Mode (Production)", 2)
    add_para(doc, 'Server mode is the primary version, designed for 24/7 production monitoring.')
    add_para(doc, 'Exclusive features:', bold=True)
    for item in [
        'Cybersecurity dark mode web dashboard (glassmorphism, neon, gradients)',
        'Interactive charts: risk distribution (doughnut) and leak sources (bars)',
        'Triage table with filters for risk, status, category, country, and language',
        'Real-time progress bar with percentage during scans',
        'Automatic scheduling via APScheduler (default: 08:00, 14:00, 22:00 BRT)',
        'Telegram alerts: system startup, scan start, and summary with new links',
        'SSE (Server-Sent Events) for real-time dashboard push notifications',
        'Complete REST API for external integrations',
        'Systemd service for continuous operation without login (Linux)',
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "7.2 Desktop Mode (Ad-hoc Query)", 2)
    add_para(doc, 'Desktop mode is a standalone GUI built with Tkinter for on-demand queries.')
    for item in [
        'Dark mode graphical interface with progress bar and metrics',
        'Result export to CSV and JSON',
        'Independent operation \u2014 no web server needed',
        'Compilable to standalone .exe via cx_Freeze',
    ]:
        add_bullet(doc, item)

    add_heading_styled(doc, "08. Step-by-Step Installation", 1)
    add_heading_styled(doc, "8.1 System Requirements", 2)
    add_pro_table(doc,
        ["Requirement", "Minimum", "Recommended"],
        [
            ["Python", "3.10", "3.12+"],
            ["Operating System", "Linux or Windows", "Ubuntu 22.04 LTS (Server)"],
            ["RAM", "512 MB", "1 GB"],
            ["Disk", "100 MB", "500 MB (database can grow)"],
            ["Network", "Internet access", "No firewall restrictions for HTTPS"],
        ],
        col_widths=[4, 5, 8],
    )
    add_heading_styled(doc, "8.2 Step 1: Clone the Repository", 2)
    add_code(doc, "git clone <repository-url>\ncd Doc-Tracker")
    add_heading_styled(doc, "8.3 Step 2: Create Virtual Environment", 2)
    add_para(doc, 'Linux / macOS:')
    add_code(doc, "python3 -m venv venv\nsource venv/bin/activate")
    add_para(doc, 'Windows:')
    add_code(doc, "python -m venv venv\nvenv\\Scripts\\activate")
    add_heading_styled(doc, "8.4 Step 3: Install Dependencies", 2)
    add_code(doc, "pip install -r requirements.txt")
    add_heading_styled(doc, "8.5 Step 4: Configure Environment Variables", 2)
    add_code(doc, "cp .env.example .env\nnano .env   # edit with actual values")
    add_heading_styled(doc, "8.6 Step 5: Run", 2)
    add_para(doc, 'Choose the appropriate execution mode (Section 10).')
    add_heading_styled(doc, "8.7 Step 6 (Optional): Install as Service", 2)
    add_para(doc, 'For production Linux servers, install as a systemd service:')
    add_code(doc,
        "# Edit the service file to adjust paths and user:\n"
        "nano server/osint-dlp.service\n\n"
        "# Copy and activate:\n"
        "sudo cp server/osint-dlp.service /etc/systemd/system/\n"
        "sudo systemctl daemon-reload\n"
        "sudo systemctl enable osint-dlp\n"
        "sudo systemctl start osint-dlp"
    )

    add_heading_styled(doc, "09. Environment Variables Configuration", 1)
    add_pro_table(doc,
        ["Variable", "Description", "Default", "Required"],
        [
            ["DATABASE_URL", "SQLite connection URL", "sqlite+aiosqlite:///osint_dlp.db", "Yes"],
            ["API_HOST", "Host the server listens on", "0.0.0.0", "No"],
            ["API_PORT", "Web server port", "8443", "No"],
            ["TELEGRAM_BOT_TOKEN", "Telegram bot token (from @BotFather)", "\u2014", "For alerts"],
            ["TELEGRAM_CHAT_ID", "Telegram chat/group ID for alerts", "\u2014", "For alerts"],
            ["SCAN_SCHEDULE_HOURS", "Auto-scan hours (comma-separated)", "8,14,22", "No"],
            ["SERVER_MODE", "Enable server features (scheduler, alerts)", "true", "No"],
            ["DASHBOARD_URL", "Public dashboard URL (used in Telegram links)", "http://localhost:8443", "No"],
            ["PROXY_URL", "SOCKS5 proxy for crawling", "\u2014", "No"],
        ],
        col_widths=[4, 5.5, 4.5, 3],
    )

    add_heading_styled(doc, "10. Execution and Commands", 1)
    add_pro_table(doc,
        ["Mode", "Command", "Description"],
        [
            ["Local Development", "python run.py", "Local web server without scheduler or Telegram"],
            ["Server (Production)", "python server/run_server.py", "Dashboard + Telegram + Auto-scan 3x/day"],
            ["Desktop (GUI)", "python desktop/run_desktop.py", "Standalone graphical interface"],
            ["Build Desktop", "python setup_desktop.py build", "Generate .exe via cx_Freeze"],
            ["Data Ingestion", "python tools/ingest_anexos.py", "Process new registration PDFs"],
            ["Generate Docs", "python tools/generate_docs.py", "Recreate these documents"],
        ],
        col_widths=[4, 6, 7],
    )

    add_heading_styled(doc, "11. Web Dashboard \u2014 Visual Guide", 1)
    add_pro_table(doc,
        ["Section", "Description"],
        [
            ["Header", "Logo, system status (Idle/Scanning with pulsing dot), 'Start Scan' button"],
            ["Progress Bar", "Appears during scans with current phase, percentage, and dork/URL detail"],
            ["Metrics", "5 cards with Total Findings, Critical, High, Medium, and Low (hover animation)"],
            ["Risk Chart", "Doughnut chart showing distribution by risk level"],
            ["Leak Sources", "Horizontal bar chart with platforms containing the most leaks"],
            ["Triage Table", "Complete finding list with 13 filterable columns"],
            ["Detail Panel", "Slide-in side panel with complete finding information"],
            ["Actions", "Buttons to classify: Investigate, False Positive, Resolved, Notified"],
            ["Pagination", "Navigation between result pages (20 per page)"],
            ["Toasts", "Floating notifications for action confirmations and new alerts"],
        ],
        col_widths=[4, 13],
    )

    add_heading_styled(doc, "12. Telegram Alerts", 1)
    add_pro_table(doc,
        ["Alert", "When", "Content"],
        [
            ["System Online", "Service startup", "Dashboard URL, timestamp, configured scan schedule"],
            ["Scan Started", "Start of each scan", "Number of dorks to execute, timestamp"],
            ["Scan Summary", "End of each scan", "Count by level + list of ALL new links grouped by risk"],
        ],
        col_widths=[3.5, 4, 9.5],
    )
    add_heading_styled(doc, "12.2 Incremental Behavior", 2)
    add_para(doc,
        'The system uses URL-based deduplication: if a document was found in a previous scan, '
        'it is not reprocessed or notified again. This ensures only NEW leaks generate alerts.'
    )

    add_heading_styled(doc, "13. REST API \u2014 Complete Reference", 1)
    add_pro_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/", "Complete HTML dashboard"],
            ["GET", "/api/dashboard", "Aggregated metrics: totals, by risk, platform, category, last scan"],
            ["GET", "/api/findings?page=N&risk_level=X&status=Y&category=Z&country=C&language=L", "List findings with pagination and filters"],
            ["PATCH", "/api/findings/{id}/status", "Update status: pending, investigating, false_positive, resolved, notified"],
            ["DELETE", "/api/findings/{id}", "Soft delete (marks is_deleted=True, keeps in database)"],
            ["GET", "/api/scans", "History of last 50 scans"],
            ["POST", "/api/scans/trigger", "Trigger manual scan (rejects if one is already running)"],
            ["GET", "/api/scans/progress", "Current progress: phase, current, total, detail"],
            ["GET", "/api/stream", "SSE stream with real-time new_finding events"],
        ],
        col_widths=[1.5, 8, 7.5],
    )

    add_heading_styled(doc, "14. Security and Privacy", 1)
    add_pro_table(doc,
        ["Measure", "Description"],
        [
            ["In-Memory Download", "Documents are downloaded into BytesIO and processed without touching the filesystem"],
            ["PII Masking", "CPFs, CNPJs, and emails are masked in snippets (e.g., ***.456.789-**)"],
            ["Excluded Domains", "Official sites (timacagro.com.br, phosphea.com, roullier.com) are never scanned"],
            ["Soft Delete", "Data is never physically removed \u2014 marked as inactive for audit purposes"],
            ["Protected .env", "Sensitive credentials are kept out of the repository via .gitignore"],
            ["Anti-block Delays", "Randomized intervals (2-4s) between search requests"],
            ["URL Deduplication", "Each URL is processed only once across the entire history"],
        ],
        col_widths=[4, 13],
    )

    add_heading_styled(doc, "15. Maintenance and Data Ingestion", 1)
    add_para(doc, 'To add new companies or update registration data:', bold=True)
    add_para(doc, '1. Place registration PDF files in the anexos/ folder (at project root).')
    add_para(doc, '2. Run the ingestion script:')
    add_code(doc, "python tools/ingest_anexos.py")
    add_para(doc, '3. The script automatically extracts:')
    for item in [
        'Company name and main CNPJ',
        'State registration, founding date, share capital',
        'Partners with participation percentage',
        'Administrators with CPFs',
        'Branches with CNPJs and addresses',
        'All CNPJs and CPFs mentioned in the document',
        'Emails, phone numbers, suppliers, and banking references',
    ]:
        add_bullet(doc, item)
    add_para(doc, '4. Data is saved to data/entities.json and original PDFs are deleted.')
    add_para(doc, '5. The next scan will automatically use the new entities.')

    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_e = p_end.add_run("_______________________________________________")
    run_e.font.color.rgb = CYAN
    run_e.bold = True
    p_end2 = doc.add_paragraph()
    p_end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end2.paragraph_format.space_before = Pt(16)
    run_e2 = p_end2.add_run("Confidential Document\nRoullier Group Security Team\nVersion 2.1 \u2014 2026")
    run_e2.font.size = Pt(10)
    run_e2.font.color.rgb = TEXT_LIGHT
    run_e2.italic = True

    return doc


def docx_to_pdf_native(docx_path, pdf_path):
    doc_in = Document(str(docx_path))
    pdf = fitz.open()
    page_w, page_h = fitz.paper_size("A4")
    margin = 56
    usable_w = page_w - 2 * margin

    all_lines = []
    for para in doc_in.paragraphs:
        text = para.text.strip()
        if not text:
            all_lines.append(("blank", "", 8))
            continue

        is_cover_title = False
        is_cover_sub = False
        is_heading1 = False
        is_heading2 = False
        is_heading3 = False
        is_code = False
        is_bullet = False

        style_name = para.style.name if para.style else ""

        if any(r.bold and r.font.size and r.font.size >= Pt(28) for r in para.runs if r.text.strip()):
            is_cover_title = True
        elif any(r.font.size and r.font.size >= Pt(16) for r in para.runs if r.text.strip()):
            is_cover_sub = True
        elif "Heading 1" in style_name:
            is_heading1 = True
        elif "Heading 2" in style_name:
            is_heading2 = True
        elif "Heading 3" in style_name:
            is_heading3 = True
        elif "No Spacing" in style_name or any(r.font.name == "Consolas" for r in para.runs):
            is_code = True
        elif "Bullet" in style_name or "List" in style_name:
            is_bullet = True

        if is_cover_title:
            all_lines.append(("cover_title", text, 40))
        elif is_cover_sub:
            all_lines.append(("cover_sub", text, 24))
        elif is_heading1:
            all_lines.append(("h1", text, 26))
        elif is_heading2:
            all_lines.append(("h2", text, 20))
        elif is_heading3:
            all_lines.append(("h3", text, 18))
        elif is_code:
            for cl in text.split("\n"):
                all_lines.append(("code", cl, 13))
        elif is_bullet:
            all_lines.append(("bullet", text, 15))
        elif text.startswith("____"):
            all_lines.append(("line", text, 10))
        else:
            all_lines.append(("body", text, 15))

    for table in doc_in.tables:
        all_lines.append(("blank", "", 4))
        for r_idx, row in enumerate(table.rows):
            row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
            if r_idx == 0:
                all_lines.append(("table_header", row_text, 14))
            else:
                all_lines.append(("table_row", row_text, 13))
        all_lines.append(("blank", "", 4))

    page = pdf.new_page(width=page_w, height=page_h)
    y = margin
    is_on_cover = True

    for kind, text, height in all_lines:
        if kind == "cover_title" and is_on_cover:
            y = page_h * 0.30
            page.insert_text((margin, y), text, fontsize=28, fontname="hebo", color=(0.12, 0.25, 0.69))
            y += 10
            rect = fitz.Rect(margin, y, page_w - margin, y + 2)
            page.draw_rect(rect, color=(0.02, 0.84, 0.63), fill=(0.02, 0.84, 0.63))
            y += 20
            continue
        elif kind == "cover_sub" and is_on_cover:
            page.insert_text((margin, y), text, fontsize=13, fontname="helv", color=(0.39, 0.45, 0.55))
            y += height
            continue
        elif kind == "blank" and is_on_cover and y > page_h * 0.5:
            page = pdf.new_page(width=page_w, height=page_h)
            y = margin
            is_on_cover = False
            continue

        if y + height > page_h - margin:
            page = pdf.new_page(width=page_w, height=page_h)
            y = margin

        if kind == "blank":
            y += height
        elif kind == "line":
            rect = fitz.Rect(margin, y - 2, page_w - margin, y)
            page.draw_rect(rect, color=(0.02, 0.84, 0.63), fill=(0.02, 0.84, 0.63))
            y += 10
        elif kind == "h1":
            y += 10
            page.insert_text((margin, y), text, fontsize=16, fontname="hebo", color=(0.12, 0.25, 0.69))
            y += 6
            rect = fitz.Rect(margin, y, margin + 80, y + 2)
            page.draw_rect(rect, color=(0.02, 0.84, 0.63), fill=(0.02, 0.84, 0.63))
            y += height
        elif kind == "h2":
            y += 6
            page.insert_text((margin, y), text, fontsize=13, fontname="hebo", color=(0.12, 0.25, 0.69))
            y += height
        elif kind == "h3":
            y += 4
            page.insert_text((margin, y), text, fontsize=11, fontname="hebo", color=(0.23, 0.35, 0.55))
            y += height
        elif kind == "code":
            bg = fitz.Rect(margin + 10, y - 10, page_w - margin, y + 4)
            page.draw_rect(bg, color=None, fill=(0.94, 0.96, 0.98))
            page.insert_text((margin + 14, y), text, fontsize=8.5, fontname="cour", color=(0.12, 0.16, 0.23))
            y += 13
        elif kind == "bullet":
            page.insert_text((margin + 10, y), "\u2022", fontsize=10, fontname="helv", color=(0.02, 0.84, 0.63))
            words = text.split()
            cx = margin + 24
            for word in words:
                tw = fitz.get_text_length(word + " ", fontsize=10, fontname="helv")
                if cx + tw > page_w - margin:
                    y += 13
                    cx = margin + 24
                    if y > page_h - margin:
                        page = pdf.new_page(width=page_w, height=page_h)
                        y = margin
                page.insert_text((cx, y), word + " ", fontsize=10, fontname="helv", color=(0.12, 0.16, 0.22))
                cx += tw
            y += height
        elif kind == "table_header":
            bg = fitz.Rect(margin, y - 10, page_w - margin, y + 4)
            page.draw_rect(bg, color=None, fill=(0.06, 0.09, 0.16))
            page.insert_text((margin + 4, y), text[:120], fontsize=8, fontname="hebo", color=(1, 1, 1))
            y += height
        elif kind == "table_row":
            if y + 14 > page_h - margin:
                page = pdf.new_page(width=page_w, height=page_h)
                y = margin
            page.insert_text((margin + 4, y), text[:140], fontsize=8, fontname="helv", color=(0.12, 0.16, 0.22))
            y += height
        else:
            words = text.split()
            cx = margin
            for word in words:
                tw = fitz.get_text_length(word + " ", fontsize=10, fontname="helv")
                if cx + tw > page_w - margin:
                    y += 14
                    cx = margin
                    if y > page_h - margin:
                        page = pdf.new_page(width=page_w, height=page_h)
                        y = margin
                page.insert_text((cx, y), word + " ", fontsize=10, fontname="helv", color=(0.12, 0.16, 0.22))
                cx += tw
            y += height

    for i in range(len(pdf)):
        pg = pdf[i]
        pg.insert_text(
            (page_w - margin - 30, page_h - 20),
            f"{i + 1}/{len(pdf)}",
            fontsize=8, fontname="helv", color=(0.58, 0.64, 0.72)
        )

    pdf.save(str(pdf_path))
    pdf.close()


def main():
    print("[*] Gerando documentacao PT-BR (DOCX)...")
    doc_pt = build_pt()
    pt_docx = DOCS_DIR / "OSINT_DLP_Documentacao_PT-BR.docx"
    doc_pt.save(str(pt_docx))
    print(f"    -> {pt_docx.name}")

    print("[*] Gerando documentacao EN (DOCX)...")
    doc_en = build_en()
    en_docx = DOCS_DIR / "OSINT_DLP_Documentation_EN.docx"
    doc_en.save(str(en_docx))
    print(f"    -> {en_docx.name}")

    print("[*] Convertendo para PDF (PT-BR)...")
    pt_pdf = DOCS_DIR / "OSINT_DLP_Documentacao_PT-BR.pdf"
    docx_to_pdf_native(pt_docx, pt_pdf)
    print(f"    -> {pt_pdf.name}")

    print("[*] Convertendo para PDF (EN)...")
    en_pdf = DOCS_DIR / "OSINT_DLP_Documentation_EN.pdf"
    docx_to_pdf_native(en_docx, en_pdf)
    print(f"    -> {en_pdf.name}")

    print("[+] 4 documentos gerados com sucesso!")


if __name__ == "__main__":
    main()
