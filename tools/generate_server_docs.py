import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import fitz

BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "documentacao" / "servidor"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

BLUE_MID = RGBColor(0x1E, 0x40, 0xAF)
CYAN = RGBColor(0x06, 0xD6, 0xA0)
TEXT_DARK = RGBColor(0x1F, 0x29, 0x37)
TEXT_MED = RGBColor(0x47, 0x55, 0x69)
TEXT_LIGHT = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TBL_HDR = "0F172A"
TBL_ALT = "F1F5F9"
TBL_BRD = "CBD5E1"


def shade(cell, c):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}"/>'))

def brd(cell, **kw):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kw.items():
        borders.append(parse_xml(f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{val}"/>'))
    tcPr.append(borders)

def h(doc, text, level=1):
    hd = doc.add_heading(level=level)
    run = hd.add_run(text)
    run.font.color.rgb = BLUE_MID
    run.font.size = Pt({1:18,2:14,3:12}.get(level,12))
    sb, sa = {1:(24,8),2:(16,6),3:(12,4)}.get(level,(12,4))
    hd.paragraph_format.space_before = Pt(sb)
    hd.paragraph_format.space_after = Pt(sa)
    if level == 1:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        p2.add_run().add_break()
    return hd

def p(doc, text, bold=False, size=11, color=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or TEXT_DARK
    run.bold = bold
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.line_spacing = Pt(16)
    return par

def bullet(doc, text):
    par = doc.add_paragraph(style="List Bullet")
    par.clear()
    run = par.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(1.5)

def code(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.left_indent = Cm(1)
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    par._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>'))

def tbl(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(ht)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = WHITE
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(c, TBL_HDR); brd(c, top=TBL_BRD, bottom=TBL_BRD, left=TBL_BRD, right=TBL_BRD)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5); run.font.color.rgb = TEXT_DARK
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 1: shade(c, TBL_ALT)
            brd(c, top=TBL_BRD, bottom=TBL_BRD, left=TBL_BRD, right=TBL_BRD)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t

def cover(doc, subtitle, version, footer):
    for _ in range(6): doc.add_paragraph().paragraph_format.space_after = Pt(0)
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run("_______________________________________________"); lr.font.color.rgb = CYAN; lr.font.size = Pt(14); lr.bold = True
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; tp.paragraph_format.space_before = Pt(20)
    tr = tp.add_run("OSINT & DLP"); tr.font.size = Pt(36); tr.bold = True; tr.font.color.rgb = BLUE_MID
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp.paragraph_format.space_before = Pt(4)
    sr = sp.add_run("WEB SERVER"); sr.font.size = Pt(18); sr.font.color.rgb = CYAN; sr.font.name = "Calibri Light"
    l2 = doc.add_paragraph(); l2.alignment = WD_ALIGN_PARAGRAPH.CENTER; l2.paragraph_format.space_before = Pt(8)
    l2r = l2.add_run("_______________________________________________"); l2r.font.color.rgb = CYAN; l2r.font.size = Pt(14); l2r.bold = True
    for _ in range(2): doc.add_paragraph().paragraph_format.space_after = Pt(0)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subr = sub.add_run(subtitle); subr.font.size = Pt(16); subr.font.color.rgb = TEXT_MED
    ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER; ver.paragraph_format.space_before = Pt(12)
    vr = ver.add_run(version); vr.font.size = Pt(11); vr.font.color.rgb = TEXT_LIGHT
    for _ in range(6): doc.add_paragraph().paragraph_format.space_after = Pt(0)
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(footer); fr.font.size = Pt(9); fr.font.color.rgb = TEXT_LIGHT; fr.italic = True
    doc.add_page_break()

def end_page(doc, text):
    doc.add_page_break()
    for _ in range(8): doc.add_paragraph().paragraph_format.space_after = Pt(0)
    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = ep.add_run("_______________________________________________"); er.font.color.rgb = CYAN; er.bold = True
    ep2 = doc.add_paragraph(); ep2.alignment = WD_ALIGN_PARAGRAPH.CENTER; ep2.paragraph_format.space_before = Pt(16)
    er2 = ep2.add_run(text); er2.font.size = Pt(10); er2.font.color.rgb = TEXT_LIGHT; er2.italic = True

def init_doc():
    doc = Document()
    sty = doc.styles["Normal"]; sty.font.name = "Calibri"; sty.font.size = Pt(11); sty.font.color.rgb = TEXT_DARK
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    return doc

def add_toc(doc, title, items):
    h(doc, title, 1)
    for num, t in items:
        par = doc.add_paragraph()
        run = par.add_run(f"{num}   {t}"); run.font.size = Pt(11); run.font.color.rgb = TEXT_DARK
        par.paragraph_format.space_after = Pt(4); par.paragraph_format.left_indent = Cm(1)
    doc.add_page_break()


def build_pt():
    doc = init_doc()
    cover(doc, "Grupo Roullier \u2014 Guia do Servidor Web", "Vers\u00e3o 2.1  |  Confidencial \u2014 Uso Interno", "Grupo Roullier Security Team \u00b7 2026")
    add_toc(doc, "Sum\u00e1rio", [
        ("01", "O Que \u00e9 o Modo Servidor"),
        ("02", "Arquitetura do Servidor"),
        ("03", "Requisitos do Sistema"),
        ("04", "Instala\u00e7\u00e3o Passo a Passo"),
        ("05", "Configura\u00e7\u00e3o de Vari\u00e1veis"),
        ("06", "Dashboard Web \u2014 Guia Visual"),
        ("07", "Sistema de Scans"),
        ("08", "Agendamento Autom\u00e1tico"),
        ("09", "API REST \u2014 Refer\u00eancia Completa"),
        ("10", "Gerenciamento de Findings"),
        ("11", "Instala\u00e7\u00e3o como Servi\u00e7o (systemd)"),
        ("12", "Monitoramento e Logs"),
        ("13", "Seguran\u00e7a e Privacidade"),
    ])

    h(doc, "01. O Que \u00e9 o Modo Servidor", 1)
    p(doc,
        'O modo Servidor do OSINT & DLP System \u00e9 a vers\u00e3o principal para ambientes de produ\u00e7\u00e3o. '
        'Ele executa um dashboard web completo com interface de cybersecurity (dark mode, glassmorphism, gr\u00e1ficos interativos), '
        'API REST para integra\u00e7\u00f5es externas, e agendamento autom\u00e1tico de scans.'
    )
    p(doc, 'Funcionalidades principais:', bold=True)
    for item in [
        'Dashboard web cybersecurity com m\u00e9tricas em tempo real',
        'Gr\u00e1ficos interativos (distribui\u00e7\u00e3o de risco e fontes de vazamento)',
        'Tabela de triagem com filtros avan\u00e7ados (risco, status, categoria, pa\u00eds, idioma)',
        'Barra de progresso em tempo real durante scans',
        'Agendamento autom\u00e1tico via APScheduler (3x/dia configur\u00e1vel)',
        'API REST completa para integra\u00e7\u00f5es com ferramentas de SIEM/SOAR',
        'SSE (Server-Sent Events) para push notifications no dashboard',
        'Banco de dados SQLite persistente com hist\u00f3rico completo',
        'Deduplica\u00e7\u00e3o por URL \u2014 s\u00f3 processa documentos novos',
        'Execu\u00e7\u00e3o como servi\u00e7o systemd para opera\u00e7\u00e3o 24/7',
    ]:
        bullet(doc, item)

    h(doc, "02. Arquitetura do Servidor", 1)
    p(doc, 'O servidor \u00e9 constru\u00eddo com FastAPI + Uvicorn e utiliza os m\u00f3dulos compartilhados do core:')
    tbl(doc, ["Componente", "Tecnologia", "Fun\u00e7\u00e3o"], [
        ["Web Framework", "FastAPI 0.135+", "Roteamento, valida\u00e7\u00e3o, middleware CORS"],
        ["Servidor HTTP", "Uvicorn", "ASGI server com hot-reload"],
        ["Templates", "Jinja2", "Renderiza\u00e7\u00e3o do HTML do dashboard"],
        ["Gr\u00e1ficos", "Chart.js 4", "Doughnut e bar charts interativos"],
        ["Tipografia", "Google Fonts (Inter, JetBrains Mono)", "Fontes profissionais"],
        ["Banco de Dados", "SQLAlchemy Async + aiosqlite", "ORM ass\u00edncrono com SQLite"],
        ["Agendamento", "APScheduler (AsyncIO)", "Scans autom\u00e1ticos 3x/dia"],
        ["Eventos", "SSE (Server-Sent Events)", "Notifica\u00e7\u00f5es push em tempo real"],
    ], col_widths=[4, 5, 8])

    h(doc, "2.1 Arquivos do Servidor", 2)
    tbl(doc, ["Arquivo", "Fun\u00e7\u00e3o"], [
        ["server/run_server.py", "Entry point: configura scheduler, lifespan e inicia uvicorn"],
        ["server/osint-dlp.service", "Unit file systemd para rodar como servi\u00e7o"],
        ["api/app.py", "Cria a aplica\u00e7\u00e3o FastAPI com CORS, rotas e lifespan"],
        ["api/routes/dashboard.py", "Rota GET / (HTML) e GET /api/dashboard (m\u00e9tricas JSON)"],
        ["api/routes/findings.py", "CRUD de findings (listagem, filtros, status, soft delete)"],
        ["api/routes/scans.py", "Trigger, progresso e hist\u00f3rico de scans"],
        ["api/routes/stream.py", "SSE endpoint para eventos em tempo real"],
        ["api/templates/index.html", "Template HTML do dashboard (Jinja2)"],
        ["static/css/dashboard.css", "461 linhas de CSS dark mode cybersecurity"],
        ["static/js/dashboard.js", "308 linhas de Javascript (polling, charts, tabela, toasts)"],
    ], col_widths=[5, 12])

    h(doc, "03. Requisitos do Sistema", 1)
    tbl(doc, ["Requisito", "M\u00ednimo", "Recomendado"], [
        ["Python", "3.10", "3.12+"],
        ["Sistema Operacional", "Linux ou Windows", "Ubuntu 22.04 LTS"],
        ["RAM", "512 MB", "1 GB"],
        ["Disco", "100 MB + banco SQLite", "500 MB"],
        ["Rede", "Porta 8443 aberta", "Sem firewall para HTTPS sa\u00edda"],
        ["Navegador (acesso)", "Chrome, Firefox, Edge", "Qualquer navegador moderno"],
    ], col_widths=[4, 5, 8])

    h(doc, "04. Instala\u00e7\u00e3o Passo a Passo", 1)
    p(doc, '1. Clonar o reposit\u00f3rio:')
    code(doc, "git clone <url-do-repositorio>\ncd Doc-Tracker")
    p(doc, '2. Criar ambiente virtual:')
    code(doc, "python3 -m venv venv\nsource venv/bin/activate")
    p(doc, '3. Instalar depend\u00eancias:')
    code(doc, "pip install -r requirements.txt")
    p(doc, '4. Configurar vari\u00e1veis de ambiente:')
    code(doc, "cp .env.example .env\nnano .env")
    p(doc, '5. Executar o servidor:')
    code(doc, "python server/run_server.py")
    p(doc, '6. Abrir o dashboard no navegador:')
    code(doc, "http://<ip-do-servidor>:8443")

    h(doc, "05. Configura\u00e7\u00e3o de Vari\u00e1veis", 1)
    tbl(doc, ["Vari\u00e1vel", "Descri\u00e7\u00e3o", "Padr\u00e3o", "Obrigat\u00f3ria"], [
        ["DATABASE_URL", "URL de conex\u00e3o SQLite", "sqlite+aiosqlite:///osint_dlp.db", "Sim"],
        ["API_HOST", "Host em que o servidor escuta", "0.0.0.0", "N\u00e3o"],
        ["API_PORT", "Porta do servidor web", "8443", "N\u00e3o"],
        ["SCAN_SCHEDULE_HOURS", "Hor\u00e1rios de auto-scan (separados por v\u00edrgula)", "8,14,22", "N\u00e3o"],
        ["SERVER_MODE", "Ativar funcionalidades de servidor", "true", "N\u00e3o"],
        ["DASHBOARD_URL", "URL p\u00fablica do dashboard", "http://localhost:8443", "N\u00e3o"],
        ["PROXY_URL", "Proxy SOCKS5 para crawling", "(vazio)", "N\u00e3o"],
    ], col_widths=[4, 5.5, 4.5, 3])

    h(doc, "06. Dashboard Web \u2014 Guia Visual", 1)
    p(doc,
        'O dashboard \u00e9 a interface principal do servidor. Segue padr\u00f5es de design de ferramentas profissionais '
        'de cybersecurity: tema dark mode (#0a0e17), glassmorphism, cores neon para status, '
        'tipografia Inter e JetBrains Mono.'
    )
    h(doc, "6.1 Tema e Cores", 2)
    tbl(doc, ["Elemento", "Cor", "Hex"], [
        ["Fundo principal", "Dark Navy", "#0A0E17"],
        ["Cards e pain\u00e9is", "Dark Blue", "#111827"],
        ["Bordas e inputs", "Slate", "#1E293B"],
        ["Texto prim\u00e1rio", "Branco quente", "#F1F5F9"],
        ["Texto secund\u00e1rio", "Cinza claro", "#94A3B8"],
        ["Acento principal", "Verde neon", "#06D6A0"],
        ["Acento secund\u00e1rio", "Azul", "#3B82F6"],
        ["Risco Cr\u00edtico", "Vermelho", "#FF3B5C"],
        ["Risco Alto", "Laranja", "#FF9F43"],
        ["Risco M\u00e9dio", "Amarelo", "#FECA57"],
        ["Risco Baixo", "Verde", "#06D6A0"],
    ], col_widths=[4, 4, 9])

    h(doc, "6.2 Se\u00e7\u00f5es do Dashboard", 2)
    tbl(doc, ["Se\u00e7\u00e3o", "Descri\u00e7\u00e3o"], [
        ["Header", "Logo 'OSINT & DLP', indicador de status (Idle/Scanning com dot pulsante), bot\u00e3o 'Iniciar Scan'"],
        ["Barra de Progresso", "Aparece durante scans: fase atual (Crawling/Inspecting), percentual, detalhe do dork/URL sendo processado"],
        ["Cards de M\u00e9tricas", "5 cards com Total, Cr\u00edtico, Alto, M\u00e9dio e Baixo \u2014 efeito hover com scale transform"],
        ["Gr\u00e1fico de Risco", "Doughnut chart (Chart.js) mostrando propor\u00e7\u00e3o por n\u00edvel de risco"],
        ["Fontes de Vazamento", "Bar chart horizontal com as plataformas que mais possuem vazamentos"],
        ["Tabela de Triagem", "Lista completa de findings com 13 colunas, filtros no topo, ordena\u00e7\u00e3o por score desc"],
        ["Painel de Detalhes", "Slide-in lateral ao clicar em uma linha \u2014 mostra snippets mascarados, metadados, a\u00e7\u00f5es"],
        ["Toast Notifications", "Notifica\u00e7\u00f5es flutuantes para confirmar a\u00e7\u00f5es (status atualizado, scan conclu\u00eddo)"],
        ["Pagina\u00e7\u00e3o", "Navega\u00e7\u00e3o Previous/Next \u2014 20 findings por p\u00e1gina"],
    ], col_widths=[4, 13])

    h(doc, "6.3 Colunas da Tabela de Triagem", 2)
    tbl(doc, ["Coluna", "Descri\u00e7\u00e3o"], [
        ["Data", "Data/hora de detec\u00e7\u00e3o do finding (formato DD/MM/YYYY HH:MM)"],
        ["Score", "Pontua\u00e7\u00e3o num\u00e9rica de risco (1-100), fundo colorido por n\u00edvel"],
        ["Pa\u00eds", "Bandeira + c\u00f3digo do pa\u00eds de origem (BR, US, FR, PT)"],
        ["Entidade", "Nome da entidade do grupo encontrada"],
        ["Fonte", "Plataforma de origem (scribd, github, slideshare, web)"],
        ["Tipo", "Formato do arquivo (PDF, DOCX, XLSX, HTML, TXT)"],
        ["Autor", "Autor extra\u00eddo dos metadados do documento"],
        ["CPFs", "Quantidade de CPFs detectados"],
        ["CNPJs", "Quantidade de CNPJs detectados"],
        ["Categoria", "Classifica\u00e7\u00e3o funcional (Financeiro, RH, TI, Corporativo)"],
        ["Status", "Estado atual (Pending, Investigating, Resolved, Notified, False Positive)"],
        ["T\u00edtulo", "T\u00edtulo do documento/p\u00e1gina web"],
        ["A\u00e7\u00f5es", "Bot\u00f5es de a\u00e7\u00e3o r\u00e1pida (alterar status)"],
    ], col_widths=[3, 14])

    h(doc, "6.4 A\u00e7\u00f5es de Triagem", 2)
    tbl(doc, ["A\u00e7\u00e3o", "Descri\u00e7\u00e3o"], [
        ["Investigating", "Marca o finding como em investiga\u00e7\u00e3o ativa"],
        ["False Positive", "Marca como falso positivo (n\u00e3o \u00e9 um vazamento real)"],
        ["Resolved", "Marca como resolvido (a\u00e7\u00e3o tomada)"],
        ["Notified", "Marca como notificado (equipe ou fornecedor avisado)"],
        ["Delete (soft)", "Remove visualmente sem apagar do banco \u2014 mant\u00e9m para auditoria"],
    ], col_widths=[4, 13])

    h(doc, "07. Sistema de Scans", 1)
    h(doc, "7.1 Scan Manual", 2)
    p(doc, 'Clique no bot\u00e3o "Iniciar Scan" no header do dashboard. O scan roda em background e:')
    for item in [
        'Gera ~80 dorks com base na Intelligence Matrix',
        'Executa cada dork no DuckDuckGo com delays de 2-4 segundos',
        'Coleta todas as URLs e remove duplicatas do banco de dados',
        'Baixa, extrai texto, inspeciona e classifica cada documento',
        'Insere findings novos no banco em tempo real',
        'Atualiza m\u00e9tricas e gr\u00e1ficos automaticamente',
    ]:
        bullet(doc, item)
    p(doc, 'Dura\u00e7\u00e3o t\u00edpica: 5-15 minutos dependendo do n\u00famero de resultados.')

    h(doc, "7.2 Progresso em Tempo Real", 2)
    tbl(doc, ["Fase", "%", "Descri\u00e7\u00e3o"], [
        ["Starting", "0%", "Inicializando m\u00f3dulos e gerando dorks"],
        ["Crawling", "0-50%", "Executando dorks no motor de busca"],
        ["Inspecting", "50-99%", "Baixando e analisando cada URL"],
        ["Completed", "100%", "Scan finalizado \u2014 findings salvos no banco"],
    ], col_widths=[3, 2, 12])
    p(doc, 'O frontend faz polling a cada 2 segundos em /api/scans/progress para atualizar a barra.')

    h(doc, "08. Agendamento Autom\u00e1tico", 1)
    p(doc,
        'O APScheduler (AsyncIO) executa scans automaticamente nos hor\u00e1rios configurados na '
        'vari\u00e1vel SCAN_SCHEDULE_HOURS. O padr\u00e3o \u00e9 3x/dia: 08:00, 14:00 e 22:00 (hor\u00e1rio BRT).'
    )
    tbl(doc, ["Configura\u00e7\u00e3o", "Valor", "Resultado"], [
        ["SCAN_SCHEDULE_HOURS=8,14,22", "Padr\u00e3o", "Scans \u00e0s 08:00, 14:00 e 22:00"],
        ["SCAN_SCHEDULE_HOURS=6,12,18,23", "Personalizado", "4 scans/dia"],
        ["SCAN_SCHEDULE_HOURS=0", "Desativado", "Apenas scans manuais"],
    ], col_widths=[6, 3, 8])
    p(doc, 'O scheduler inicia junto com o servidor e roda dentro do event loop do FastAPI (lifespan).')

    h(doc, "09. API REST \u2014 Refer\u00eancia Completa", 1)
    h(doc, "9.1 Endpoints", 2)
    tbl(doc, ["M\u00e9todo", "Endpoint", "Descri\u00e7\u00e3o"], [
        ["GET", "/", "Dashboard HTML completo"],
        ["GET", "/api/dashboard", "M\u00e9tricas: totais, por risco, por plataforma, por categoria, \u00faltimo scan"],
        ["GET", "/api/findings?page=N", "Listar findings (20/p\u00e1gina)"],
        ["GET", "/api/findings?risk_level=critical", "Filtrar por risco"],
        ["GET", "/api/findings?status=pending", "Filtrar por status"],
        ["GET", "/api/findings?category=financeiro", "Filtrar por categoria"],
        ["GET", "/api/findings?country=BR", "Filtrar por pa\u00eds"],
        ["GET", "/api/findings?language=pt", "Filtrar por idioma"],
        ["PATCH", "/api/findings/{id}/status", "Atualizar status (body: {status: 'investigating'})"],
        ["DELETE", "/api/findings/{id}", "Soft delete (is_deleted=True)"],
        ["GET", "/api/scans", "\u00daltimos 50 scans"],
        ["POST", "/api/scans/trigger", "Disparar scan manual"],
        ["GET", "/api/scans/progress", "Progresso: phase, current, total, detail"],
        ["GET", "/api/stream", "SSE stream de new_finding"],
    ], col_widths=[1.5, 6, 9.5])

    h(doc, "9.2 Exemplo de Resposta: /api/dashboard", 2)
    code(doc,
        '{\n'
        '  "total_findings": 98,\n'
        '  "by_risk": {"critical": 5, "high": 12, "medium": 31, "low": 50},\n'
        '  "by_platform": {"scribd": 15, "web": 40, ...},\n'
        '  "by_category": {"financeiro": 20, "rh": 8, ...},\n'
        '  "last_scan": "2026-04-14T08:00:15Z"\n'
        '}'
    )

    h(doc, "10. Gerenciamento de Findings", 1)
    h(doc, "10.1 Modelo de Dados: Finding", 2)
    tbl(doc, ["Campo", "Tipo", "Descri\u00e7\u00e3o"], [
        ["id", "Integer (PK)", "Identificador \u00fanico"],
        ["scan_id", "Integer (FK)", "Refer\u00eancia ao scan que o encontrou"],
        ["url", "String", "URL completa do documento"],
        ["title", "String", "T\u00edtulo da p\u00e1gina/documento"],
        ["source_platform", "String", "Plataforma de origem"],
        ["file_type", "String", "Formato: pdf, docx, xlsx, html, txt"],
        ["risk_level", "String", "critical, high, medium, low"],
        ["risk_score", "Integer", "Score num\u00e9rico 1-100"],
        ["category", "String", "Categoria funcional"],
        ["entity_matched", "String", "Entidades do grupo encontradas"],
        ["cpf_count", "Integer", "Quantidade de CPFs"],
        ["cnpj_count", "Integer", "Quantidade de CNPJs"],
        ["financial_count", "Integer", "Termos financeiros"],
        ["author", "String", "Autor do documento (metadados)"],
        ["country", "String", "Pa\u00eds de origem (TLD)"],
        ["language", "String", "Idioma detectado"],
        ["status", "String", "pending, investigating, resolved, false_positive, notified"],
        ["is_deleted", "Boolean", "Soft delete flag"],
        ["created_at", "DateTime", "Data de detec\u00e7\u00e3o"],
    ], col_widths=[3.5, 3, 10.5])

    h(doc, "10.2 Soft Delete", 2)
    p(doc,
        'Ao deletar um finding via API ou dashboard, o registro N\u00c3O \u00e9 removido fisicamente. '
        'Apenas o campo is_deleted \u00e9 marcado como True. Isso garante que o hist\u00f3rico completo '
        'est\u00e1 dispon\u00edvel para auditoria forense. Consultas padr\u00e3o excluem registros deletados.'
    )

    h(doc, "11. Instala\u00e7\u00e3o como Servi\u00e7o (systemd)", 1)
    p(doc, 'Para manter o servidor rodando 24/7 sem login, instale como servi\u00e7o systemd:')
    p(doc, '1. Editar o arquivo de servi\u00e7o para ajustar caminhos e usu\u00e1rio:')
    code(doc, "nano server/osint-dlp.service")
    p(doc, '2. Copiar para systemd e ativar:')
    code(doc, "sudo cp server/osint-dlp.service /etc/systemd/system/\nsudo systemctl daemon-reload\nsudo systemctl enable osint-dlp\nsudo systemctl start osint-dlp")
    p(doc, '3. Verificar status:')
    code(doc, "sudo systemctl status osint-dlp")
    p(doc, '4. Ver logs em tempo real:')
    code(doc, "sudo journalctl -u osint-dlp -f")
    h(doc, "11.1 Comandos de Gerenciamento", 2)
    tbl(doc, ["Comando", "A\u00e7\u00e3o"], [
        ["sudo systemctl start osint-dlp", "Iniciar o servi\u00e7o"],
        ["sudo systemctl stop osint-dlp", "Parar o servi\u00e7o"],
        ["sudo systemctl restart osint-dlp", "Reiniciar (aplica atualiza\u00e7\u00f5es)"],
        ["sudo systemctl status osint-dlp", "Ver status e PID"],
        ["sudo systemctl enable osint-dlp", "Habilitar auto-start no boot"],
        ["sudo systemctl disable osint-dlp", "Desabilitar auto-start"],
        ["sudo journalctl -u osint-dlp -f", "Logs em tempo real"],
        ["sudo journalctl -u osint-dlp --since today", "Logs de hoje"],
    ], col_widths=[7, 10])

    h(doc, "12. Monitoramento e Logs", 1)
    p(doc, 'O servidor gera logs estruturados via Python logging:')
    code(doc, "2026-04-14 08:00:01 | INFO     | __main__ | OSINT/DLP Server Mode - Port 8443\n2026-04-14 08:00:01 | INFO     | __main__ | Schedule: 8,14,22\n2026-04-14 08:00:01 | INFO     | __main__ | Scheduler started\n2026-04-14 08:00:15 | INFO     | api.routes.scans | === OSINT Scan Started ===\n2026-04-14 08:05:32 | INFO     | api.routes.scans | === Scan Complete: 98 findings from 142 URLs ===")
    p(doc, 'Indicadores de sa\u00fade:', bold=True)
    tbl(doc, ["Indicador", "Como Verificar", "Saud\u00e1vel"], [
        ["Servidor online", "curl http://localhost:8443/api/dashboard", "Retorna JSON com m\u00e9tricas"],
        ["Scheduler ativo", "Logs: 'Scheduler started'", "3 jobs adicionados"],
        ["Scans rodando", "GET /api/scans/progress", "phase != 'failed'"],
        ["Banco acess\u00edvel", "GET /api/findings", "Retorna lista (mesmo vazia)"],
        ["Disco", "du -sh osint_dlp.db", "Tamanho crescente ap\u00f3s scans"],
    ], col_widths=[3.5, 6, 7.5])

    h(doc, "13. Seguran\u00e7a e Privacidade", 1)
    tbl(doc, ["Medida", "Descri\u00e7\u00e3o"], [
        ["Download em Mem\u00f3ria", "Documentos baixados em BytesIO, processados sem gravar em disco"],
        ["PII Mascarados", "CPFs, CNPJs e emails mascarados nos snippets exibidos"],
        ["Dom\u00ednios Exclu\u00eddos", "Sites oficiais do grupo nunca s\u00e3o varridos (timacagro.com.br, etc.)"],
        ["Soft Delete", "Dados nunca removidos fisicamente \u2014 mant\u00e9m hist\u00f3rico de auditoria"],
        [".env Protegido", "Credenciais sens\u00edveis fora do reposit\u00f3rio Git"],
        ["CORS Configurado", "Middleware CORS permite acesso controlado ao dashboard"],
        ["Delays Anti-bloqueio", "2-4 segundos entre requisi\u00e7\u00f5es de busca"],
        ["Dedup por URL", "Cada URL processada uma \u00fanica vez em todo o hist\u00f3rico"],
        ["Restart Autom\u00e1tico", "Servi\u00e7o systemd com Restart=always e RestartSec=10"],
    ], col_widths=[4, 13])

    end_page(doc, "Documento Confidencial\nGrupo Roullier Security Team\nVers\u00e3o 2.1 \u2014 2026")
    return doc


def build_en():
    doc = init_doc()
    cover(doc, "Roullier Group \u2014 Web Server Guide", "Version 2.1  |  Confidential \u2014 Internal Use Only", "Roullier Group Security Team \u00b7 2026")
    add_toc(doc, "Table of Contents", [
        ("01", "What is Server Mode"),
        ("02", "Server Architecture"),
        ("03", "System Requirements"),
        ("04", "Step-by-Step Installation"),
        ("05", "Environment Configuration"),
        ("06", "Web Dashboard \u2014 Visual Guide"),
        ("07", "Scan System"),
        ("08", "Automatic Scheduling"),
        ("09", "REST API \u2014 Complete Reference"),
        ("10", "Finding Management"),
        ("11", "Service Installation (systemd)"),
        ("12", "Monitoring and Logs"),
        ("13", "Security and Privacy"),
    ])

    h(doc, "01. What is Server Mode", 1)
    p(doc,
        'The Server mode of the OSINT & DLP System is the primary version for production environments. '
        'It runs a full web dashboard with a cybersecurity interface (dark mode, glassmorphism, interactive charts), '
        'REST API for external integrations, and automatic scan scheduling.'
    )
    p(doc, 'Key features:', bold=True)
    for item in [
        'Cybersecurity web dashboard with real-time metrics',
        'Interactive charts (risk distribution and leak sources)',
        'Triage table with advanced filters (risk, status, category, country, language)',
        'Real-time progress bar during scans',
        'Automatic scheduling via APScheduler (3x/day configurable)',
        'Complete REST API for SIEM/SOAR tool integrations',
        'SSE (Server-Sent Events) for dashboard push notifications',
        'Persistent SQLite database with complete history',
        'URL-based deduplication \u2014 only processes new documents',
        'Systemd service for 24/7 operation',
    ]:
        bullet(doc, item)

    h(doc, "02. Server Architecture", 1)
    p(doc, 'The server is built with FastAPI + Uvicorn and uses the shared core modules:')
    tbl(doc, ["Component", "Technology", "Function"], [
        ["Web Framework", "FastAPI 0.135+", "Routing, validation, CORS middleware"],
        ["HTTP Server", "Uvicorn", "ASGI server with hot-reload"],
        ["Templates", "Jinja2", "Dashboard HTML rendering"],
        ["Charts", "Chart.js 4", "Interactive doughnut and bar charts"],
        ["Typography", "Google Fonts (Inter, JetBrains Mono)", "Professional fonts"],
        ["Database", "SQLAlchemy Async + aiosqlite", "Async ORM with SQLite"],
        ["Scheduling", "APScheduler (AsyncIO)", "Automatic scans 3x/day"],
        ["Events", "SSE (Server-Sent Events)", "Real-time push notifications"],
    ], col_widths=[4, 5, 8])

    h(doc, "2.1 Server Files", 2)
    tbl(doc, ["File", "Function"], [
        ["server/run_server.py", "Entry point: configures scheduler, lifespan, and starts uvicorn"],
        ["server/osint-dlp.service", "Systemd unit file for service operation"],
        ["api/app.py", "Creates FastAPI application with CORS, routes, and lifespan"],
        ["api/routes/dashboard.py", "GET / (HTML) and GET /api/dashboard (JSON metrics)"],
        ["api/routes/findings.py", "Finding CRUD (listing, filters, status, soft delete)"],
        ["api/routes/scans.py", "Trigger, progress, and scan history"],
        ["api/routes/stream.py", "SSE endpoint for real-time events"],
        ["api/templates/index.html", "Dashboard HTML template (Jinja2)"],
        ["static/css/dashboard.css", "461 lines of dark mode cybersecurity CSS"],
        ["static/js/dashboard.js", "308 lines of Javascript (polling, charts, table, toasts)"],
    ], col_widths=[5, 12])

    h(doc, "03. System Requirements", 1)
    tbl(doc, ["Requirement", "Minimum", "Recommended"], [
        ["Python", "3.10", "3.12+"],
        ["Operating System", "Linux or Windows", "Ubuntu 22.04 LTS"],
        ["RAM", "512 MB", "1 GB"],
        ["Disk", "100 MB + SQLite database", "500 MB"],
        ["Network", "Port 8443 open", "No firewall for outbound HTTPS"],
        ["Browser (access)", "Chrome, Firefox, Edge", "Any modern browser"],
    ], col_widths=[4, 5, 8])

    h(doc, "04. Step-by-Step Installation", 1)
    p(doc, '1. Clone the repository:')
    code(doc, "git clone <repository-url>\ncd Doc-Tracker")
    p(doc, '2. Create virtual environment:')
    code(doc, "python3 -m venv venv\nsource venv/bin/activate")
    p(doc, '3. Install dependencies:')
    code(doc, "pip install -r requirements.txt")
    p(doc, '4. Configure environment variables:')
    code(doc, "cp .env.example .env\nnano .env")
    p(doc, '5. Run the server:')
    code(doc, "python server/run_server.py")
    p(doc, '6. Open the dashboard in browser:')
    code(doc, "http://<server-ip>:8443")

    h(doc, "05. Environment Configuration", 1)
    tbl(doc, ["Variable", "Description", "Default", "Required"], [
        ["DATABASE_URL", "SQLite connection URL", "sqlite+aiosqlite:///osint_dlp.db", "Yes"],
        ["API_HOST", "Host the server listens on", "0.0.0.0", "No"],
        ["API_PORT", "Web server port", "8443", "No"],
        ["SCAN_SCHEDULE_HOURS", "Auto-scan hours (comma-separated)", "8,14,22", "No"],
        ["SERVER_MODE", "Enable server features", "true", "No"],
        ["DASHBOARD_URL", "Public dashboard URL", "http://localhost:8443", "No"],
        ["PROXY_URL", "SOCKS5 proxy for crawling", "(empty)", "No"],
    ], col_widths=[4, 5.5, 4.5, 3])

    h(doc, "06. Web Dashboard \u2014 Visual Guide", 1)
    p(doc,
        'The dashboard is the main server interface. It follows professional cybersecurity tool design: '
        'dark mode theme (#0a0e17), glassmorphism, neon status colors, Inter and JetBrains Mono typography.'
    )
    h(doc, "6.1 Dashboard Sections", 2)
    tbl(doc, ["Section", "Description"], [
        ["Header", "Logo 'OSINT & DLP', status indicator (Idle/Scanning with pulsing dot), 'Start Scan' button"],
        ["Progress Bar", "Appears during scans: current phase (Crawling/Inspecting), percentage, dork/URL detail"],
        ["Metric Cards", "5 cards: Total, Critical, High, Medium, Low \u2014 hover effect with scale transform"],
        ["Risk Chart", "Doughnut chart (Chart.js) showing proportion by risk level"],
        ["Leak Sources", "Horizontal bar chart with platforms containing the most leaks"],
        ["Triage Table", "Complete finding list with 13 columns, top filters, sorted by score desc"],
        ["Detail Panel", "Slide-in panel on row click \u2014 shows masked snippets, metadata, actions"],
        ["Toast Notifications", "Floating notifications confirming actions (status updated, scan complete)"],
        ["Pagination", "Previous/Next navigation \u2014 20 findings per page"],
    ], col_widths=[4, 13])

    h(doc, "6.2 Triage Table Columns", 2)
    tbl(doc, ["Column", "Description"], [
        ["Date", "Detection date/time (DD/MM/YYYY HH:MM format)"],
        ["Score", "Numeric risk score (1-100), colored background by level"],
        ["Country", "Flag + country code (BR, US, FR, PT)"],
        ["Entity", "Group entity name found"],
        ["Source", "Origin platform (scribd, github, slideshare, web)"],
        ["Type", "File format (PDF, DOCX, XLSX, HTML, TXT)"],
        ["Author", "Author from document metadata"],
        ["CPFs", "Number of CPFs detected"],
        ["CNPJs", "Number of CNPJs detected"],
        ["Category", "Functional classification (Financial, HR, IT, Corporate)"],
        ["Status", "Current state (Pending, Investigating, Resolved, Notified, False Positive)"],
        ["Title", "Document/web page title"],
        ["Actions", "Quick action buttons (change status)"],
    ], col_widths=[3, 14])

    h(doc, "6.3 Triage Actions", 2)
    tbl(doc, ["Action", "Description"], [
        ["Investigating", "Marks the finding as under active investigation"],
        ["False Positive", "Marks as false positive (not a real leak)"],
        ["Resolved", "Marks as resolved (action taken)"],
        ["Notified", "Marks as notified (team or vendor alerted)"],
        ["Delete (soft)", "Visually removes without database deletion \u2014 kept for audit"],
    ], col_widths=[4, 13])

    h(doc, "07. Scan System", 1)
    h(doc, "7.1 Manual Scan", 2)
    p(doc, 'Click "Start Scan" in the dashboard header. The scan runs in background and:')
    for item in [
        'Generates ~80 dorks based on the Intelligence Matrix',
        'Executes each dork on DuckDuckGo with 2-4 second delays',
        'Collects all URLs and removes duplicates from database',
        'Downloads, extracts text, inspects, and classifies each document',
        'Inserts new findings into database in real-time',
        'Automatically updates metrics and charts',
    ]:
        bullet(doc, item)
    p(doc, 'Typical duration: 5-15 minutes depending on results.')

    h(doc, "7.2 Real-Time Progress", 2)
    tbl(doc, ["Phase", "%", "Description"], [
        ["Starting", "0%", "Initializing modules and generating dorks"],
        ["Crawling", "0-50%", "Executing dorks on search engine"],
        ["Inspecting", "50-99%", "Downloading and analyzing each URL"],
        ["Completed", "100%", "Scan finished \u2014 findings saved to database"],
    ], col_widths=[3, 2, 12])

    h(doc, "08. Automatic Scheduling", 1)
    p(doc, 'APScheduler (AsyncIO) runs scans at hours configured in SCAN_SCHEDULE_HOURS. Default: 3x/day at 08:00, 14:00, 22:00 BRT.')
    tbl(doc, ["Configuration", "Value", "Result"], [
        ["SCAN_SCHEDULE_HOURS=8,14,22", "Default", "Scans at 08:00, 14:00, 22:00"],
        ["SCAN_SCHEDULE_HOURS=6,12,18,23", "Custom", "4 scans/day"],
        ["SCAN_SCHEDULE_HOURS=0", "Disabled", "Manual scans only"],
    ], col_widths=[6, 3, 8])

    h(doc, "09. REST API \u2014 Complete Reference", 1)
    tbl(doc, ["Method", "Endpoint", "Description"], [
        ["GET", "/", "Complete HTML dashboard"],
        ["GET", "/api/dashboard", "Metrics: totals, by risk, platform, category, last scan"],
        ["GET", "/api/findings?page=N", "List findings (20/page)"],
        ["GET", "/api/findings?risk_level=critical", "Filter by risk"],
        ["GET", "/api/findings?status=pending", "Filter by status"],
        ["GET", "/api/findings?category=financeiro", "Filter by category"],
        ["GET", "/api/findings?country=BR", "Filter by country"],
        ["GET", "/api/findings?language=pt", "Filter by language"],
        ["PATCH", "/api/findings/{id}/status", "Update status (body: {status: 'investigating'})"],
        ["DELETE", "/api/findings/{id}", "Soft delete (is_deleted=True)"],
        ["GET", "/api/scans", "Last 50 scans"],
        ["POST", "/api/scans/trigger", "Trigger manual scan"],
        ["GET", "/api/scans/progress", "Progress: phase, current, total, detail"],
        ["GET", "/api/stream", "SSE stream of new_finding events"],
    ], col_widths=[1.5, 6, 9.5])

    h(doc, "10. Finding Management", 1)
    h(doc, "10.1 Data Model: Finding", 2)
    tbl(doc, ["Field", "Type", "Description"], [
        ["id", "Integer (PK)", "Unique identifier"],
        ["scan_id", "Integer (FK)", "Reference to the scan that found it"],
        ["url", "String", "Full document URL"],
        ["title", "String", "Page/document title"],
        ["risk_level", "String", "critical, high, medium, low"],
        ["risk_score", "Integer", "Numeric score 1-100"],
        ["category", "String", "Functional category"],
        ["entity_matched", "String", "Group entities found"],
        ["cpf_count", "Integer", "Number of CPFs"],
        ["cnpj_count", "Integer", "Number of CNPJs"],
        ["country", "String", "Origin country (TLD)"],
        ["language", "String", "Detected language"],
        ["status", "String", "pending, investigating, resolved, false_positive, notified"],
        ["is_deleted", "Boolean", "Soft delete flag"],
        ["created_at", "DateTime", "Detection date"],
    ], col_widths=[3.5, 3, 10.5])

    h(doc, "11. Service Installation (systemd)", 1)
    p(doc, 'To keep the server running 24/7 without login, install as a systemd service:')
    code(doc, "# Edit service file for your paths and user:\nnano server/osint-dlp.service\n\n# Copy and activate:\nsudo cp server/osint-dlp.service /etc/systemd/system/\nsudo systemctl daemon-reload\nsudo systemctl enable osint-dlp\nsudo systemctl start osint-dlp")
    tbl(doc, ["Command", "Action"], [
        ["sudo systemctl start osint-dlp", "Start the service"],
        ["sudo systemctl stop osint-dlp", "Stop the service"],
        ["sudo systemctl restart osint-dlp", "Restart (apply updates)"],
        ["sudo systemctl status osint-dlp", "View status and PID"],
        ["sudo journalctl -u osint-dlp -f", "Real-time logs"],
    ], col_widths=[7, 10])

    h(doc, "12. Monitoring and Logs", 1)
    tbl(doc, ["Indicator", "How to Check", "Healthy"], [
        ["Server online", "curl http://localhost:8443/api/dashboard", "Returns JSON with metrics"],
        ["Scheduler active", "Logs: 'Scheduler started'", "3 jobs added"],
        ["Scans running", "GET /api/scans/progress", "phase != 'failed'"],
        ["DB accessible", "GET /api/findings", "Returns list (even if empty)"],
        ["Disk", "du -sh osint_dlp.db", "Growing size after scans"],
    ], col_widths=[3.5, 6, 7.5])

    h(doc, "13. Security and Privacy", 1)
    tbl(doc, ["Measure", "Description"], [
        ["In-Memory Download", "Documents downloaded into BytesIO, processed without disk writes"],
        ["PII Masked", "CPFs, CNPJs, emails masked in displayed snippets"],
        ["Excluded Domains", "Official group sites never scanned (timacagro.com.br, etc.)"],
        ["Soft Delete", "Data never physically removed \u2014 maintains audit history"],
        ["Protected .env", "Sensitive credentials kept out of Git repository"],
        ["CORS Configured", "CORS middleware controls dashboard access"],
        ["Anti-block Delays", "2-4 seconds between search requests"],
        ["URL Dedup", "Each URL processed only once across all history"],
        ["Auto-Restart", "Systemd service with Restart=always and RestartSec=10"],
    ], col_widths=[4, 13])

    end_page(doc, "Confidential Document\nRoullier Group Security Team\nVersion 2.1 \u2014 2026")
    return doc


def docx_to_pdf(docx_path, pdf_path):
    doc_in = Document(str(docx_path))
    pdf = fitz.open()
    pw, ph = fitz.paper_size("A4")
    m = 56
    lines = []
    for para in doc_in.paragraphs:
        t = para.text.strip()
        if not t: lines.append(("blank","",8)); continue
        sn = para.style.name if para.style else ""
        if any(r.bold and r.font.size and r.font.size >= Pt(28) for r in para.runs if r.text.strip()): lines.append(("ct",t,40))
        elif any(r.font.size and r.font.size >= Pt(16) for r in para.runs if r.text.strip()): lines.append(("cs",t,24))
        elif "Heading 1" in sn: lines.append(("h1",t,26))
        elif "Heading 2" in sn: lines.append(("h2",t,20))
        elif "Heading 3" in sn: lines.append(("h3",t,18))
        elif "No Spacing" in sn or any(r.font.name=="Consolas" for r in para.runs):
            for cl in t.split("\n"): lines.append(("code",cl,13))
        elif "Bullet" in sn or "List" in sn: lines.append(("bullet",t,15))
        elif t.startswith("____"): lines.append(("line",t,10))
        else: lines.append(("body",t,15))
    for tl in doc_in.tables:
        lines.append(("blank","",4))
        for ri, row in enumerate(tl.rows):
            rt = " | ".join(c.text.strip().replace("\n"," ") for c in row.cells)
            lines.append(("th" if ri==0 else "tr", rt, 14 if ri==0 else 13))
        lines.append(("blank","",4))
    page = pdf.new_page(width=pw, height=ph); y=m; oc=True
    for kind,text,ht in lines:
        if kind=="ct" and oc:
            y=ph*0.30; page.insert_text((m,y),text,fontsize=28,fontname="hebo",color=(0.12,0.25,0.69)); y+=10
            page.draw_rect(fitz.Rect(m,y,pw-m,y+2),color=(0.02,0.84,0.63),fill=(0.02,0.84,0.63)); y+=20; continue
        elif kind=="cs" and oc: page.insert_text((m,y),text,fontsize=13,fontname="helv",color=(0.39,0.45,0.55)); y+=ht; continue
        elif kind=="blank" and oc and y>ph*0.5: page=pdf.new_page(width=pw,height=ph); y=m; oc=False; continue
        if y+ht>ph-m: page=pdf.new_page(width=pw,height=ph); y=m
        if kind=="blank": y+=ht
        elif kind=="line": page.draw_rect(fitz.Rect(m,y-2,pw-m,y),color=(0.02,0.84,0.63),fill=(0.02,0.84,0.63)); y+=10
        elif kind=="h1":
            y+=10; page.insert_text((m,y),text,fontsize=16,fontname="hebo",color=(0.12,0.25,0.69)); y+=6
            page.draw_rect(fitz.Rect(m,y,m+80,y+2),color=(0.02,0.84,0.63),fill=(0.02,0.84,0.63)); y+=ht
        elif kind=="h2": y+=6; page.insert_text((m,y),text,fontsize=13,fontname="hebo",color=(0.12,0.25,0.69)); y+=ht
        elif kind=="h3": y+=4; page.insert_text((m,y),text,fontsize=11,fontname="hebo",color=(0.23,0.35,0.55)); y+=ht
        elif kind=="code":
            page.draw_rect(fitz.Rect(m+10,y-10,pw-m,y+4),color=None,fill=(0.94,0.96,0.98))
            page.insert_text((m+14,y),text,fontsize=8.5,fontname="cour",color=(0.12,0.16,0.23)); y+=13
        elif kind=="bullet":
            page.insert_text((m+10,y),"\u2022",fontsize=10,fontname="helv",color=(0.02,0.84,0.63))
            words=text.split(); cx=m+24
            for w in words:
                tw=fitz.get_text_length(w+" ",fontsize=10,fontname="helv")
                if cx+tw>pw-m: y+=13; cx=m+24
                if y>ph-m: page=pdf.new_page(width=pw,height=ph); y=m
                page.insert_text((cx,y),w+" ",fontsize=10,fontname="helv",color=(0.12,0.16,0.22)); cx+=tw
            y+=ht
        elif kind=="th":
            page.draw_rect(fitz.Rect(m,y-10,pw-m,y+4),color=None,fill=(0.06,0.09,0.16))
            page.insert_text((m+4,y),text[:120],fontsize=8,fontname="hebo",color=(1,1,1)); y+=ht
        elif kind=="tr":
            if y+14>ph-m: page=pdf.new_page(width=pw,height=ph); y=m
            page.insert_text((m+4,y),text[:140],fontsize=8,fontname="helv",color=(0.12,0.16,0.22)); y+=ht
        else:
            words=text.split(); cx=m
            for w in words:
                tw=fitz.get_text_length(w+" ",fontsize=10,fontname="helv")
                if cx+tw>pw-m: y+=14; cx=m
                if y>ph-m: page=pdf.new_page(width=pw,height=ph); y=m
                page.insert_text((cx,y),w+" ",fontsize=10,fontname="helv",color=(0.12,0.16,0.22)); cx+=tw
            y+=ht
    for i in range(len(pdf)):
        pdf[i].insert_text((pw-m-30,ph-20),f"{i+1}/{len(pdf)}",fontsize=8,fontname="helv",color=(0.58,0.64,0.72))
    pdf.save(str(pdf_path)); pdf.close()


def main():
    print("[*] Server docs PT-BR (DOCX)...")
    pt = build_pt(); f1 = DOCS_DIR/"OSINT_DLP_Servidor_PT-BR.docx"; pt.save(str(f1)); print(f"    -> {f1.name}")
    print("[*] Server docs EN (DOCX)...")
    en = build_en(); f2 = DOCS_DIR/"OSINT_DLP_Server_EN.docx"; en.save(str(f2)); print(f"    -> {f2.name}")
    print("[*] PDF PT-BR...")
    f3 = DOCS_DIR/"OSINT_DLP_Servidor_PT-BR.pdf"; docx_to_pdf(f1, f3); print(f"    -> {f3.name}")
    print("[*] PDF EN...")
    f4 = DOCS_DIR/"OSINT_DLP_Server_EN.pdf"; docx_to_pdf(f2, f4); print(f"    -> {f4.name}")
    print("[+] 4 documentos servidor gerados com sucesso!")

if __name__ == "__main__":
    main()
