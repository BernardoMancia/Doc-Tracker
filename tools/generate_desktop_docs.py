import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import fitz

BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "documentacao" / "desktop"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

BLUE_MID = RGBColor(0x1E, 0x40, 0xAF)
CYAN = RGBColor(0x06, 0xD6, 0xA0)
TEXT_DARK = RGBColor(0x1F, 0x29, 0x37)
TEXT_MED = RGBColor(0x47, 0x55, 0x69)
TEXT_LIGHT = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TBL_HDR = "0F172A"
TBL_ALT = "F1F5F9"
TBL_BRD = "CBD5E1"


def shade(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    )


def border(cell, **kw):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kw.items():
        borders.append(parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{val}"/>'
        ))
    tcPr.append(borders)


def h(doc, text, level=1):
    hd = doc.add_heading(level=level)
    run = hd.add_run(text)
    run.font.color.rgb = BLUE_MID
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    spaces = {1: (24, 8), 2: (16, 6), 3: (12, 4)}
    sb, sa = spaces.get(level, (12, 4))
    hd.paragraph_format.space_before = Pt(sb)
    hd.paragraph_format.space_after = Pt(sa)
    if level == 1:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run()
        r2.add_break()
    return hd


def p(doc, text, bold=False, size=11, color=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or TEXT_DARK
    run.bold = bold
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.line_spacing = Pt(16)
    return par


def bullet(doc, text):
    par = doc.add_paragraph(style="List Bullet")
    par.clear()
    run = par.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(1.5)
    return par


def code(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.left_indent = Cm(1)
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    par._p.get_or_add_pPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
    )
    return par


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        pr = c.paragraphs[0]
        run = pr.add_run(ht)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(c, TBL_HDR)
        border(c, top=TBL_BRD, bottom=TBL_BRD, left=TBL_BRD, right=TBL_BRD)
    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            pr = c.paragraphs[0]
            run = pr.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT_DARK
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 1:
                shade(c, TBL_ALT)
            border(c, top=TBL_BRD, bottom=TBL_BRD, left=TBL_BRD, right=TBL_BRD)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t


def cover(doc, subtitle, version, footer):
    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run("_______________________________________________")
    lr.font.color.rgb = CYAN
    lr.font.size = Pt(14)
    lr.bold = True

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(20)
    tr = tp.add_run("OSINT & DLP")
    tr.font.size = Pt(36)
    tr.bold = True
    tr.font.color.rgb = BLUE_MID

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(4)
    sr = sp.add_run("DESKTOP APPLICATION")
    sr.font.size = Pt(18)
    sr.font.color.rgb = CYAN
    sr.font.name = "Calibri Light"

    l2p = doc.add_paragraph()
    l2p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    l2p.paragraph_format.space_before = Pt(8)
    l2r = l2p.add_run("_______________________________________________")
    l2r.font.color.rgb = CYAN
    l2r.font.size = Pt(14)
    l2r.bold = True

    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subr = sub.add_run(subtitle)
    subr.font.size = Pt(16)
    subr.font.color.rgb = TEXT_MED
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.paragraph_format.space_before = Pt(12)
    vr = ver.add_run(version)
    vr.font.size = Pt(11)
    vr.font.color.rgb = TEXT_LIGHT

    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(footer)
    fr.font.size = Pt(9)
    fr.font.color.rgb = TEXT_LIGHT
    fr.italic = True
    doc.add_page_break()


def end_page(doc, text):
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = ep.add_run("_______________________________________________")
    er.font.color.rgb = CYAN
    er.bold = True
    ep2 = doc.add_paragraph()
    ep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep2.paragraph_format.space_before = Pt(16)
    er2 = ep2.add_run(text)
    er2.font.size = Pt(10)
    er2.font.color.rgb = TEXT_LIGHT
    er2.italic = True


def build_pt():
    doc = Document()
    sty = doc.styles["Normal"]
    sty.font.name = "Calibri"
    sty.font.size = Pt(11)
    sty.font.color.rgb = TEXT_DARK
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    cover(doc,
        "Grupo Roullier \u2014 Guia Completo da Vers\u00e3o Desktop",
        "Vers\u00e3o 2.1  |  Confidencial \u2014 Uso Interno",
        "Grupo Roullier Security Team \u00b7 2026",
    )

    h(doc, "Sum\u00e1rio", 1)
    toc = [
        ("01", "O Que \u00e9 a Vers\u00e3o Desktop"),
        ("02", "Desktop vs. Server \u2014 Diferen\u00e7as"),
        ("03", "Requisitos do Sistema"),
        ("04", "Instala\u00e7\u00e3o Passo a Passo"),
        ("05", "Interface Gr\u00e1fica \u2014 Guia Completo"),
        ("06", "Como Funciona o Scan"),
        ("07", "Tabela de Resultados \u2014 Colunas Explicadas"),
        ("08", "Painel de Detalhes"),
        ("09", "Exporta\u00e7\u00e3o de Resultados"),
        ("10", "Compila\u00e7\u00e3o do Execut\u00e1vel (.exe)"),
        ("11", "Principais Arquivos"),
        ("12", "Solu\u00e7\u00e3o de Problemas"),
        ("13", "Seguran\u00e7a e Privacidade"),
    ]
    for num, title in toc:
        par = doc.add_paragraph()
        run = par.add_run(f"{num}   {title}")
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
        par.paragraph_format.space_after = Pt(4)
        par.paragraph_format.left_indent = Cm(1)
    doc.add_page_break()

    h(doc, "01. O Que \u00e9 a Vers\u00e3o Desktop", 1)
    p(doc,
        'A vers\u00e3o Desktop do OSINT & DLP System \u00e9 uma aplica\u00e7\u00e3o gr\u00e1fica standalone (GUI) '
        'constru\u00edda com Tkinter, projetada para consultas sob demanda sem necessidade '
        'de servidor web, banco de dados persistente ou infraestrutura de rede.'
    )
    p(doc,
        'A aplica\u00e7\u00e3o executa o mesmo pipeline de intelig\u00eancia do modo Server '
        '(gera\u00e7\u00e3o de dorks, busca web, download em mem\u00f3ria, extra\u00e7\u00e3o de texto, '
        'inspe\u00e7\u00e3o regex e classifica\u00e7\u00e3o de risco), por\u00e9m em formato visual e interativo.'
    )
    p(doc, 'Principais caracter\u00edsticas:', bold=True)
    for item in [
        'Interface dark mode profissional inspirada em ferramentas de cybersecurity',
        'Execu\u00e7\u00e3o totalmente local \u2014 n\u00e3o depende de servidor web',
        'Barra de progresso em tempo real com percentual e fase atual',
        'Cards de m\u00e9tricas com contagem por n\u00edvel de risco (Cr\u00edtico/Alto/M\u00e9dio/Baixo)',
        'Tabela interativa com 12 colunas de dados',
        'Painel de detalhes expand\u00edvel ao clicar em um finding',
        'Exporta\u00e7\u00e3o para CSV com um clique',
        'Operac\u00e3o totalmente em mem\u00f3ria \u2014 zero grava\u00e7\u00e3o em disco',
        'Compil\u00e1vel para execut\u00e1vel .exe via cx_Freeze',
    ]:
        bullet(doc, item)

    h(doc, "02. Desktop vs. Server \u2014 Diferen\u00e7as", 1)
    table(doc,
        ["Caracter\u00edstica", "Desktop (GUI)", "Server (Web)"],
        [
            ["Interface", "Tkinter (aplica\u00e7\u00e3o nativa)", "Dashboard web (FastAPI + HTML)"],
            ["Banco de Dados", "Sem banco \u2014 tudo em mem\u00f3ria", "SQLite persistente"],
            ["Hist\u00f3rico de Scans", "N\u00e3o mant\u00e9m hist\u00f3rico", "Hist\u00f3rico completo com deduplica\u00e7\u00e3o"],
            ["Agendamento", "Scan manual apenas", "3x/dia autom\u00e1tico (APScheduler)"],
            ["Alertas Telegram", "N\u00e3o envia", "Startup + in\u00edcio + resumo de scan"],
            ["API REST", "N\u00e3o possui", "Endpoints completos"],
            ["SSE (Tempo Real)", "N\u00e3o possui", "Server-Sent Events"],
            ["Exporta\u00e7\u00e3o", "CSV com um clique", "Via API ou dashboard"],
            ["Deploy", "Execut\u00e1vel (.exe) ou Python direto", "Servi\u00e7o systemd no servidor"],
            ["Deduplica\u00e7\u00e3o", "Por sess\u00e3o (set de URLs)", "Por hist\u00f3rico completo (banco)"],
            ["Uso Ideal", "Consultas ad-hoc e auditorias pontuais", "Monitoramento cont\u00ednuo 24/7"],
        ],
        col_widths=[4, 6.5, 6.5],
    )

    h(doc, "03. Requisitos do Sistema", 1)
    table(doc,
        ["Requisito", "M\u00ednimo", "Recomendado"],
        [
            ["Python", "3.10 (se rodar via script)", "3.12+"],
            ["Sistema Operacional", "Windows 10 / Linux", "Windows 11"],
            ["RAM", "512 MB", "1 GB"],
            ["Disco", "50 MB (+ depend\u00eancias)", "100 MB"],
            ["Rede", "Acesso \u00e0 internet", "Sem firewall bloqueando HTTPS"],
            ["Resolu\u00e7\u00e3o M\u00ednima", "1000 x 650 px", "1300 x 780 px ou mais"],
        ],
        col_widths=[4.5, 5, 7.5],
    )
    p(doc,
        'Para rodar o execut\u00e1vel compilado (.exe), Python n\u00e3o \u00e9 necess\u00e1rio \u2014 todas as '
        'depend\u00eancias s\u00e3o empacotadas pelo cx_Freeze.'
    )

    h(doc, "04. Instala\u00e7\u00e3o Passo a Passo", 1)
    h(doc, "4.1 Op\u00e7\u00e3o A: Rodar via Python (Desenvolvedor)", 2)
    p(doc, '1. Clonar o reposit\u00f3rio:')
    code(doc, "git clone <url-do-repositorio>\ncd Doc-Tracker")
    p(doc, '2. Criar ambiente virtual:')
    p(doc, 'Windows:')
    code(doc, "python -m venv venv\nvenv\\Scripts\\activate")
    p(doc, 'Linux / macOS:')
    code(doc, "python3 -m venv venv\nsource venv/bin/activate")
    p(doc, '3. Instalar depend\u00eancias:')
    code(doc, "pip install -r requirements.txt")
    p(doc, '4. Configurar vari\u00e1veis de ambiente (opcional para desktop):')
    code(doc, "cp .env.example .env")
    p(doc,
        'Para o modo Desktop, as \u00fanicas vari\u00e1veis relevantes s\u00e3o DATABASE_URL (opcionalmente). '
        'As vari\u00e1veis de Telegram e scheduler n\u00e3o s\u00e3o usadas.'
    )
    p(doc, '5. Executar:')
    code(doc, "python desktop/run_desktop.py")

    h(doc, "4.2 Op\u00e7\u00e3o B: Rodar via Execut\u00e1vel (.exe)", 2)
    p(doc, '1. Solicitar o execut\u00e1vel compilado ao administrador do sistema.')
    p(doc, '2. Extrair a pasta OSINT_DLP_Desktop/ em qualquer local do computador.')
    p(doc, '3. Executar o arquivo OSINT_DLP_Desktop.exe diretamente (duplo clique).')
    p(doc, 'N\u00e3o \u00e9 necess\u00e1rio instalar Python nem configura\u00e7\u00f5es adicionais.')

    h(doc, "05. Interface Gr\u00e1fica \u2014 Guia Completo", 1)
    p(doc,
        'A interface segue um design dark mode profissional inspirado em ferramentas de cybersecurity. '
        'Toda a aplica\u00e7\u00e3o \u00e9 constru\u00edda em uma \u00fanica janela com as seguintes \u00e1reas:'
    )

    h(doc, "5.1 Header (Barra Superior)", 2)
    table(doc,
        ["Elemento", "Descri\u00e7\u00e3o"],
        [
            ["T\u00edtulo", "'OSINT & DLP \u2014 Consulta Desktop' em verde neon (com icone escudo)"],
            ["Bot\u00e3o Exportar CSV", "Exportar CSV \u2014 salva todos os resultados em arquivo CSV"],
            ["Bot\u00e3o Consultar", "Consultar Agora \u2014 inicia o scan (desabilita durante execu\u00e7\u00e3o)"],
        ],
        col_widths=[4, 13],
    )

    h(doc, "5.2 Barra de Progresso", 2)
    p(doc,
        'Localizada abaixo do header, mostra o andamento do scan atual. A barra verde preenche '
        'de 0% a 100%. No lado esquerdo mostra a fase ("Buscando... 5/80 dorks" ou "Analisando... 42/120 URLs"). '
        'No direito, o percentual num\u00e9rico.'
    )
    table(doc,
        ["Fase", "Percentual", "O Que Acontece"],
        [
            ["Buscando...", "0-50%", "Execu\u00e7\u00e3o de ~80 dorks no DuckDuckGo com delays anti-bloqueio"],
            ["Analisando...", "50-100%", "Download, extra\u00e7\u00e3o, inspe\u00e7\u00e3o e classifica\u00e7\u00e3o de cada URL encontrada"],
            ["Conclu\u00eddo", "100%", "Scan finalizado \u2014 total de findings exibido"],
        ],
        col_widths=[3, 3, 11],
    )

    h(doc, "5.3 Cards de M\u00e9tricas", 2)
    p(doc,
        'Cinco cartoes horizontais mostram contadores em tempo real. Cada cart\u00e3o tem um n\u00famero grande '
        '(fonte JetBrains Mono 22pt bold) e um label abaixo:'
    )
    table(doc,
        ["Card", "Cor", "Descri\u00e7\u00e3o"],
        [
            ["TOTAL", "Verde (#06D6A0)", "Soma de todos os findings detectados nesta sess\u00e3o"],
            ["CR\u00cdTICO", "Vermelho (#FF3B5C)", "Findings com score 80-100 (dados altamente sens\u00edveis)"],
            ["ALTO", "Laranja (#FF9F43)", "Findings com score 60-79 (dados sens\u00edveis moderados)"],
            ["M\u00c9DIO", "Amarelo (#FECA57)", "Findings com score 40-59 (men\u00e7\u00f5es relevantes)"],
            ["BAIXO", "Verde (#06D6A0)", "Findings com score 1-39 (men\u00e7\u00f5es gen\u00e9ricas)"],
        ],
        col_widths=[3, 4, 10],
    )
    p(doc, 'Os contadores atualizam automaticamente conforme novos findings s\u00e3o detectados durante o scan.')

    h(doc, "5.4 Status Bar (Barra Inferior)", 2)
    p(doc,
        'Na parte inferior da janela exibe o texto "Grupo Roullier \u2014 Threat Monitor Desktop" por padr\u00e3o. '
        'Ap\u00f3s um scan, mostra a data/hora da \u00faltima consulta. Ap\u00f3s exportar, mostra o caminho do arquivo salvo.'
    )

    h(doc, "06. Como Funciona o Scan", 1)
    p(doc,
        'Ao clicar em "Consultar Agora", o sistema inicia um scan completo em uma thread separada '
        'para n\u00e3o travar a interface gr\u00e1fica. O fluxo \u00e9 o seguinte:'
    )
    steps = [
        ("Inicializa\u00e7\u00e3o",
         "Carrega a Intelligence Matrix (data/entities.json) com entidades, CNPJs, CPFs, fornecedores e termos "
         "sens\u00edveis. Inicializa os m\u00f3dulos: DorkGenerator, SearchEngine, URLFilter, Downloader, TextExtractor, "
         "RegexEngine e RiskClassifier."),
        ("Gera\u00e7\u00e3o de Dorks (0-50%)",
         "O DorkGenerator cria ~80 consultas avan\u00e7adas cobrindo: raz\u00f5es sociais, CNPJs no corpo de documentos, "
         "CPFs de administradores, fornecedores, termos financeiros (faturamento, balan\u00e7o, extrato) e termos de "
         "seguran\u00e7a (senha, password, confidencial). Exclui dom\u00ednios oficiais com -site:."),
        ("Busca Web (0-50%)",
         "Cada dork \u00e9 executado no DuckDuckGo com delays randomizados de 2 a 4 segundos. "
         "A barra de progresso preenche de 0% a 50% nesta fase."),
        ("Deduplica\u00e7\u00e3o por URL",
         "URLs duplicadas encontradas em dorks diferentes s\u00e3o removidas usando um set em mem\u00f3ria."),
        ("Inspe\u00e7\u00e3o de URLs (50-100%)",
         "Para cada URL \u00fanica: download em BytesIO (mem\u00f3ria), detec\u00e7\u00e3o de tipo (PDF/DOCX/XLSX/TXT), "
         "extra\u00e7\u00e3o de texto e metadados, inspe\u00e7\u00e3o por regex (CPFs, CNPJs, financeiro, entidades, termos sens\u00edveis), "
         "classifica\u00e7\u00e3o de risco (score 1-100 + n\u00edvel + categoria), detec\u00e7\u00e3o de pa\u00eds (TLD) e idioma (texto)."),
        ("Exibi\u00e7\u00e3o em Tempo Real",
         "Cada finding detectado \u00e9 inserido imediatamente na tabela e os cards de m\u00e9tricas s\u00e3o atualizados. "
         "N\u00e3o \u00e9 necess\u00e1rio esperar o scan terminar para ver resultados."),
        ("Finaliza\u00e7\u00e3o",
         "Barra chega a 100%, bot\u00e3o volta a ficar ativo, status bar mostra data/hora da consulta. "
         "Os resultados ficam na mem\u00f3ria at\u00e9 fechar a aplica\u00e7\u00e3o ou iniciar novo scan."),
    ]
    for i, (title, desc) in enumerate(steps, 1):
        h(doc, f"Passo {i}: {title}", 3)
        p(doc, desc)

    h(doc, "07. Tabela de Resultados \u2014 Colunas Explicadas", 1)
    p(doc, 'A tabela central ocupa a maior parte da tela e cont\u00e9m 12 colunas:')
    table(doc,
        ["Coluna", "Descri\u00e7\u00e3o", "Exemplo"],
        [
            ["SCORE", "Pontua\u00e7\u00e3o num\u00e9rica de risco (1-100)", "85"],
            ["RISCO", "N\u00edvel textual do risco", "CR\u00cdTICO, ALTO, M\u00c9DIO, BAIXO"],
            ["PA\u00cdS", "Pa\u00eds de origem detectado pelo TLD + bandeira emoji", "[BR]"],
            ["IDIOMA", "Idioma do conte\u00fado detectado por an\u00e1lise textual", "PT, EN, ES, FR"],
            ["ENTIDADE", "Nome da entidade do grupo encontrada (primeiros 25 caracteres)", "Timac Agro Brasil"],
            ["FONTE", "Plataforma onde o documento foi encontrado", "scribd, slideshare, web"],
            ["T\u00cdTULO", "T\u00edtulo do documento/p\u00e1gina (primeiros 35 caracteres)", "Balan\u00e7o Patrimonial 2024"],
            ["TIPO", "Formato do arquivo detectado", "PDF, DOCX, XLSX, HTML"],
            ["AUTOR", "Autor extra\u00eddo dos metadados (primeiros 18 caracteres)", "Jo\u00e3o Silva"],
            ["CPFS", "Quantidade de CPFs detectados no documento", "3"],
            ["CNPJS", "Quantidade de CNPJs detectados", "1"],
            ["CATEGORIA", "Classifica\u00e7\u00e3o funcional do finding", "Financeiro, RH, TI/Seg"],
        ],
        col_widths=[2.5, 7, 7.5],
    )
    p(doc,
        'As linhas s\u00e3o coloridas conforme o n\u00edvel de risco: vermelho (Cr\u00edtico), laranja (Alto), '
        'amarelo (M\u00e9dio) e verde (Baixo). Os findings mais recentes aparecem no topo da tabela.'
    )

    h(doc, "08. Painel de Detalhes", 1)
    p(doc,
        'Ao clicar em qualquer linha da tabela, um painel de detalhes se expande abaixo mostrando '
        'informa\u00e7\u00f5es completas do finding selecionado. O painel usa fonte monoespa\u00e7ada '
        '(JetBrains Mono 10pt) para leitura t\u00e9cnica.'
    )
    p(doc, 'Informa\u00e7\u00f5es exibidas:', bold=True)
    table(doc,
        ["Campo", "Descri\u00e7\u00e3o"],
        [
            ["T\u00cdTULO", "T\u00edtulo completo do documento/p\u00e1gina"],
            ["URL", "Link completo para acesso direto ao documento"],
            ["SCORE / RISCO / CATEGORIA", "Classificaa\u00e7\u00e3o completa do finding"],
            ["ENTIDADE", "Todas as entidades do grupo encontradas (at\u00e9 5)"],
            ["PA\u00cdS / IDIOMA", "Bandeira + c\u00f3digo do pa\u00eds e idioma por extenso"],
            ["AUTOR / PUBLICADOR", "Metadados extra\u00eddos do documento (quando dispon\u00edvel)"],
            ["CPFs / CNPJs / FINANCEIROS", "Contagem detalhada de dados sens\u00edveis encontrados"],
            ["PLATAFORMA / TIPO", "Fonte e formato do arquivo"],
            ["RAZ\u00d5ES", "Lista de raz\u00f5es que levaram \u00e0 classifica\u00e7\u00e3o de risco"],
            ["TERMOS", "Termos sens\u00edveis encontrados no documento (at\u00e9 10)"],
        ],
        col_widths=[5, 12],
    )

    h(doc, "09. Exporta\u00e7\u00e3o de Resultados", 1)
    h(doc, "9.1 Exportar para CSV", 2)
    p(doc, '1. Execute um scan completo.')
    p(doc, '2. Clique no bot\u00e3o "Exportar CSV" no canto superior direito.')
    p(doc, '3. Escolha o local e nome do arquivo no di\u00e1logo de salvamento.')
    p(doc, f'4. O nome sugerido \u00e9: osint_findings_YYYYMMDD_HHMM.csv')
    p(doc, 'Colunas do CSV exportado:', bold=True)
    table(doc,
        ["Coluna CSV", "Descri\u00e7\u00e3o"],
        [
            ["risk_score", "Score num\u00e9rico (1-100)"],
            ["risk_level", "N\u00edvel: critical, high, medium, low"],
            ["country", "C\u00f3digo do pa\u00eds (BR, US, FR, etc.)"],
            ["language", "C\u00f3digo do idioma (pt, en, es, fr)"],
            ["entity_matched", "Entidades detectadas"],
            ["source_platform", "Plataforma de origem"],
            ["title", "T\u00edtulo completo"],
            ["file_type", "Formato do arquivo"],
            ["author", "Autor dos metadados"],
            ["publisher", "Publicador/criador do documento"],
            ["cpf_count", "Quantidade de CPFs"],
            ["cnpj_count", "Quantidade de CNPJs"],
            ["financial_count", "Quantidade de termos financeiros"],
            ["category", "Categoria funcional"],
            ["url", "URL completa de acesso"],
        ],
        col_widths=[4, 13],
    )

    h(doc, "10. Compila\u00e7\u00e3o do Execut\u00e1vel (.exe)", 1)
    p(doc,
        'Para distribuir a aplica\u00e7\u00e3o sem exigir Python instalado, '
        'utilize o cx_Freeze para gerar um execut\u00e1vel standalone.'
    )
    h(doc, "10.1 Pr\u00e9-requisitos", 2)
    code(doc, "pip install cx_Freeze")
    h(doc, "10.2 Compilar", 2)
    code(doc, "python setup_desktop.py build")
    p(doc,
        'O script limpa automaticamente as pastas build/ e dist/ antes de compilar. '
        'O resultado \u00e9 gerado em dist/OSINT_DLP_Desktop/.'
    )
    h(doc, "10.3 Conte\u00fado Gerado", 2)
    table(doc,
        ["Arquivo/Pasta", "Descri\u00e7\u00e3o"],
        [
            ["OSINT_DLP_Desktop.exe", "Execut\u00e1vel principal"],
            ["data/entities.json", "Base de entidades do grupo (inclu\u00edda automaticamente)"],
            ["lib/", "Bibliotecas Python empacotadas"],
            ["python3XX.dll", "Runtime Python embutido"],
        ],
        col_widths=[5, 12],
    )
    h(doc, "10.4 Pacotes Inclu\u00eddos vs. Exclu\u00eddos", 2)
    table(doc,
        ["Inclu\u00eddos", "Exclu\u00eddos (apenas do Server)"],
        [
            ["asyncio, json, csv, logging, threading", "fastapi"],
            ["tkinter, tkinter.ttk", "uvicorn, starlette"],
            ["httpx, fitz (PyMuPDF)", "jinja2"],
            ["docx (python-docx), openpyxl", "apscheduler"],
            ["ddgs, pydantic, pydantic_settings", "multipart"],
            ["sqlalchemy, aiosqlite", "test, unittest, pytest"],
        ],
        col_widths=[8.5, 8.5],
    )

    h(doc, "11. Principais Arquivos", 1)
    table(doc,
        ["Arquivo", "Linhas", "Fun\u00e7\u00e3o"],
        [
            ["desktop/app.py", "436", "Toda a l\u00f3gica da GUI: interface, scan, export, detalhes"],
            ["desktop/run_desktop.py", "10", "Entry point: resolve caminhos e chama desktop.app.main()"],
            ["setup_desktop.py", "47", "Configura\u00e7\u00e3o de build cx_Freeze para .exe"],
            ["data/entities.json", "~320", "Base de entidades (compartilhada com Server)"],
            ["config/settings.py", "~40", "Configura\u00e7\u00f5es de ambiente (carrega .env)"],
            ["config/intelligence_matrix.py", "~80", "Carrega entities.json e monta a matrix de intelig\u00eancia"],
            ["requirements.txt", "~15", "Depend\u00eancias Python (compartilhadas)"],
        ],
        col_widths=[5, 2, 10],
    )

    h(doc, "12. Solu\u00e7\u00e3o de Problemas", 1)
    table(doc,
        ["Problema", "Causa Prov\u00e1vel", "Solu\u00e7\u00e3o"],
        [
            ["Janela n\u00e3o abre", "Python sem Tkinter instalado", "Windows: reinstalar Python com op\u00e7\u00e3o tkinter. Linux: sudo apt install python3-tk"],
            ["Poucos resultados", "DuckDuckGo rate-limiting", "Esperar 5 minutos e tentar novamente. Considerar usar PROXY_URL no .env"],
            ["Erro de encoding", "Console com charset limitado", "Usar terminal com suporte UTF-8 (Windows Terminal recomendado)"],
            ["Scan demora muito", "Muitas URLs para inspecionar", "Normal \u2014 cada URL leva 1-3 segundos. ~80 dorks + delays = ~5-10 minutos total"],
            ["Bot\u00e3o CSV cinza", "Nenhum resultado na tabela", "Executar um scan completo antes de exportar"],
            [".exe n\u00e3o abre", "Antiv\u00edrus bloqueando", "Adicionar exce\u00e7\u00e3o no antiv\u00edrus para o diret\u00f3rio OSINT_DLP_Desktop"],
            ["ModuleNotFoundError", "Depend\u00eancia n\u00e3o instalada", "Executar pip install -r requirements.txt no ambiente virtual"],
            ["entities.json n\u00e3o encontrado", "Caminho relativo errado", "Executar sempre a partir da raiz do projeto (cd Doc-Tracker)"],
        ],
        col_widths=[3, 4, 10],
    )

    h(doc, "13. Seguran\u00e7a e Privacidade", 1)
    table(doc,
        ["Medida", "Descri\u00e7\u00e3o"],
        [
            ["Zero Disco", "Todos os documentos s\u00e3o baixados em BytesIO (mem\u00f3ria) e nunca gravados no filesystem"],
            ["Sem Persist\u00eancia", "Resultados existem apenas na mem\u00f3ria da aplica\u00e7\u00e3o, descartados ao fechar"],
            ["PII Mascarados", "CPFs, CNPJs e emails s\u00e3o mascarados nos snippets (ex: ***.456.789-**)"],
            ["Dom\u00ednios Exclu\u00eddos", "Sites oficiais do grupo s\u00e3o automaticamente exclu\u00eddos da varredura"],
            ["Thread Separada", "O scan roda em thread separada para n\u00e3o travar a interface"],
            ["Sem Rede Local", "A aplica\u00e7\u00e3o n\u00e3o abre portas de rede \u2014 zero superf\u00edcie de ataque"],
            ["Delays Anti-bloqueio", "Intervalos randomizados (2-4s) entre consultas para evitar rate-limiting"],
        ],
        col_widths=[4, 13],
    )

    end_page(doc, "Documento Confidencial\nGrupo Roullier Security Team\nVers\u00e3o 2.1 \u2014 2026")
    return doc


def build_en():
    doc = Document()
    sty = doc.styles["Normal"]
    sty.font.name = "Calibri"
    sty.font.size = Pt(11)
    sty.font.color.rgb = TEXT_DARK
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    cover(doc,
        "Roullier Group \u2014 Complete Desktop Application Guide",
        "Version 2.1  |  Confidential \u2014 Internal Use Only",
        "Roullier Group Security Team \u00b7 2026",
    )

    h(doc, "Table of Contents", 1)
    toc = [
        ("01", "What is the Desktop Version"),
        ("02", "Desktop vs. Server \u2014 Differences"),
        ("03", "System Requirements"),
        ("04", "Step-by-Step Installation"),
        ("05", "Graphical Interface \u2014 Complete Guide"),
        ("06", "How the Scan Works"),
        ("07", "Results Table \u2014 Columns Explained"),
        ("08", "Detail Panel"),
        ("09", "Exporting Results"),
        ("10", "Building the Executable (.exe)"),
        ("11", "Key Files"),
        ("12", "Troubleshooting"),
        ("13", "Security and Privacy"),
    ]
    for num, title in toc:
        par = doc.add_paragraph()
        run = par.add_run(f"{num}   {title}")
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
        par.paragraph_format.space_after = Pt(4)
        par.paragraph_format.left_indent = Cm(1)
    doc.add_page_break()

    h(doc, "01. What is the Desktop Version", 1)
    p(doc,
        'The Desktop version of the OSINT & DLP System is a standalone graphical application (GUI) '
        'built with Tkinter, designed for on-demand queries without the need for a web server, '
        'persistent database, or network infrastructure.'
    )
    p(doc,
        'The application runs the same intelligence pipeline as Server mode (dork generation, web search, '
        'in-memory download, text extraction, regex inspection, and risk classification), but in a visual '
        'and interactive format.'
    )
    p(doc, 'Key features:', bold=True)
    for item in [
        'Professional dark mode interface inspired by cybersecurity tools',
        'Fully local execution \u2014 no dependency on a web server',
        'Real-time progress bar with percentage and current phase',
        'Metric cards with risk-level counters (Critical/High/Medium/Low)',
        'Interactive table with 12 data columns',
        'Expandable detail panel when clicking on a finding',
        'One-click CSV export',
        'Fully in-memory operation \u2014 zero disk writes',
        'Compilable to standalone .exe via cx_Freeze',
    ]:
        bullet(doc, item)

    h(doc, "02. Desktop vs. Server \u2014 Differences", 1)
    table(doc,
        ["Feature", "Desktop (GUI)", "Server (Web)"],
        [
            ["Interface", "Tkinter (native application)", "Web dashboard (FastAPI + HTML)"],
            ["Database", "None \u2014 everything in memory", "Persistent SQLite"],
            ["Scan History", "No history kept", "Full history with deduplication"],
            ["Scheduling", "Manual scan only", "3x/day automatic (APScheduler)"],
            ["Telegram Alerts", "Not sent", "Startup + start + scan summary"],
            ["REST API", "None", "Complete endpoints"],
            ["SSE (Real-Time)", "None", "Server-Sent Events"],
            ["Export", "One-click CSV", "Via API or dashboard"],
            ["Deploy", "Executable (.exe) or direct Python", "Systemd service on server"],
            ["Deduplication", "Per session (URL set)", "Full history (database)"],
            ["Best For", "Ad-hoc queries and audits", "Continuous 24/7 monitoring"],
        ],
        col_widths=[4, 6.5, 6.5],
    )

    h(doc, "03. System Requirements", 1)
    table(doc,
        ["Requirement", "Minimum", "Recommended"],
        [
            ["Python", "3.10 (if running via script)", "3.12+"],
            ["Operating System", "Windows 10 / Linux", "Windows 11"],
            ["RAM", "512 MB", "1 GB"],
            ["Disk", "50 MB (+ dependencies)", "100 MB"],
            ["Network", "Internet access", "No firewall blocking HTTPS"],
            ["Minimum Resolution", "1000 x 650 px", "1300 x 780 px or higher"],
        ],
        col_widths=[4.5, 5, 7.5],
    )
    p(doc,
        'To run the compiled executable (.exe), Python is not required \u2014 all '
        'dependencies are bundled by cx_Freeze.'
    )

    h(doc, "04. Step-by-Step Installation", 1)
    h(doc, "4.1 Option A: Run via Python (Developer)", 2)
    p(doc, '1. Clone the repository:')
    code(doc, "git clone <repository-url>\ncd Doc-Tracker")
    p(doc, '2. Create virtual environment:')
    p(doc, 'Windows:')
    code(doc, "python -m venv venv\nvenv\\Scripts\\activate")
    p(doc, 'Linux / macOS:')
    code(doc, "python3 -m venv venv\nsource venv/bin/activate")
    p(doc, '3. Install dependencies:')
    code(doc, "pip install -r requirements.txt")
    p(doc, '4. Configure environment variables (optional for desktop):')
    code(doc, "cp .env.example .env")
    p(doc,
        'For Desktop mode, the only relevant variable is DATABASE_URL (optionally). '
        'Telegram and scheduler variables are not used.'
    )
    p(doc, '5. Run:')
    code(doc, "python desktop/run_desktop.py")

    h(doc, "4.2 Option B: Run via Executable (.exe)", 2)
    p(doc, '1. Request the compiled executable from the system administrator.')
    p(doc, '2. Extract the OSINT_DLP_Desktop/ folder to any location on the computer.')
    p(doc, '3. Run the OSINT_DLP_Desktop.exe file directly (double-click).')
    p(doc, 'No Python installation or additional configuration is required.')

    h(doc, "05. Graphical Interface \u2014 Complete Guide", 1)
    p(doc,
        'The interface follows a professional dark mode design inspired by cybersecurity tools. '
        'The entire application is built in a single window with the following areas:'
    )
    h(doc, "5.1 Header (Top Bar)", 2)
    table(doc,
        ["Element", "Description"],
        [
            ["Title", "'OSINT & DLP \u2014 Desktop Query' in neon green (with shield icon)"],
            ["Export CSV Button", "Export CSV \u2014 saves all results to a CSV file"],
            ["Query Button", "Query Now \u2014 starts the scan (disabled during execution)"],
        ],
        col_widths=[4, 13],
    )
    h(doc, "5.2 Progress Bar", 2)
    p(doc,
        'Located below the header, it shows scan progress. The green bar fills from 0% to 100%. '
        'Left side shows the phase ("Searching... 5/80 dorks" or "Analyzing... 42/120 URLs"). '
        'Right side shows the numeric percentage.'
    )
    table(doc,
        ["Phase", "Percentage", "What Happens"],
        [
            ["Searching...", "0-50%", "Execution of ~80 dorks on DuckDuckGo with anti-blocking delays"],
            ["Analyzing...", "50-100%", "Download, extraction, inspection, and classification of each URL found"],
            ["Completed", "100%", "Scan finished \u2014 total findings displayed"],
        ],
        col_widths=[3, 3, 11],
    )
    h(doc, "5.3 Metric Cards", 2)
    p(doc,
        'Five horizontal cards show real-time counters. Each card has a large number '
        '(JetBrains Mono 22pt bold) and a label below:'
    )
    table(doc,
        ["Card", "Color", "Description"],
        [
            ["TOTAL", "Green (#06D6A0)", "Sum of all findings detected in this session"],
            ["CRITICAL", "Red (#FF3B5C)", "Findings with score 80-100 (highly sensitive data)"],
            ["HIGH", "Orange (#FF9F43)", "Findings with score 60-79 (moderate sensitive data)"],
            ["MEDIUM", "Yellow (#FECA57)", "Findings with score 40-59 (relevant mentions)"],
            ["LOW", "Green (#06D6A0)", "Findings with score 1-39 (generic mentions)"],
        ],
        col_widths=[3, 4, 10],
    )

    h(doc, "06. How the Scan Works", 1)
    p(doc,
        'When clicking "Query Now", the system starts a full scan in a separate thread '
        'to keep the graphical interface responsive.'
    )
    steps = [
        ("Initialization",
         "Loads the Intelligence Matrix (data/entities.json) with entities, CNPJs, CPFs, suppliers, and "
         "sensitive terms. Initializes modules: DorkGenerator, SearchEngine, URLFilter, Downloader, "
         "TextExtractor, RegexEngine, and RiskClassifier."),
        ("Dork Generation (0-50%)",
         "DorkGenerator creates ~80 advanced queries covering: company names, CNPJs in document bodies, "
         "administrator CPFs, suppliers, financial terms (revenue, balance, statement) and security terms "
         "(password, confidential). Official domains excluded with -site:."),
        ("Web Search (0-50%)",
         "Each dork is executed on DuckDuckGo with randomized 2-4 second delays. "
         "Progress bar fills from 0% to 50% during this phase."),
        ("URL Deduplication",
         "Duplicate URLs found across different dorks are removed using an in-memory set."),
        ("URL Inspection (50-100%)",
         "For each unique URL: BytesIO download (memory), type detection (PDF/DOCX/XLSX/TXT), "
         "text and metadata extraction, regex inspection (CPFs, CNPJs, financial, entities, sensitive terms), "
         "risk classification (score 1-100 + level + category), country (TLD) and language detection (text)."),
        ("Real-Time Display",
         "Each detected finding is immediately inserted into the table and metric cards are updated. "
         "No need to wait for the scan to finish to see results."),
        ("Finalization",
         "Bar reaches 100%, button reactivates, status bar shows query date/time. "
         "Results remain in memory until the application is closed or a new scan starts."),
    ]
    for i, (title, desc) in enumerate(steps, 1):
        h(doc, f"Step {i}: {title}", 3)
        p(doc, desc)

    h(doc, "07. Results Table \u2014 Columns Explained", 1)
    table(doc,
        ["Column", "Description", "Example"],
        [
            ["SCORE", "Numeric risk score (1-100)", "85"],
            ["RISK", "Text risk level", "CRITICAL, HIGH, MEDIUM, LOW"],
            ["COUNTRY", "Origin country detected by TLD + flag emoji", "[BR]"],
            ["LANGUAGE", "Content language detected by text analysis", "PT, EN, ES, FR"],
            ["ENTITY", "Group entity found (first 25 characters)", "Timac Agro Brasil"],
            ["SOURCE", "Platform where the document was found", "scribd, slideshare, web"],
            ["TITLE", "Document/page title (first 35 characters)", "Balance Sheet 2024"],
            ["TYPE", "Detected file format", "PDF, DOCX, XLSX, HTML"],
            ["AUTHOR", "Author from metadata (first 18 characters)", "John Smith"],
            ["CPFS", "Number of CPFs detected in the document", "3"],
            ["CNPJS", "Number of CNPJs detected", "1"],
            ["CATEGORY", "Functional classification of the finding", "Financial, HR, IT/Sec"],
        ],
        col_widths=[2.5, 7, 7.5],
    )
    p(doc,
        'Rows are colored by risk level: red (Critical), orange (High), yellow (Medium), '
        'and green (Low). Most recent findings appear at the top.'
    )

    h(doc, "08. Detail Panel", 1)
    p(doc,
        'Clicking any table row expands a detail panel below showing complete finding information '
        'in monospaced font (JetBrains Mono 10pt) for technical reading.'
    )
    table(doc,
        ["Field", "Description"],
        [
            ["TITLE", "Full document/page title"],
            ["URL", "Complete access link"],
            ["SCORE / RISK / CATEGORY", "Full finding classification"],
            ["ENTITY", "All group entities found (up to 5)"],
            ["COUNTRY / LANGUAGE", "Flag + country code and language in full"],
            ["AUTHOR / PUBLISHER", "Metadata extracted from document (when available)"],
            ["CPFs / CNPJs / FINANCIAL", "Detailed count of sensitive data found"],
            ["PLATFORM / TYPE", "Source and file format"],
            ["REASONS", "List of reasons that led to the risk classification"],
            ["TERMS", "Sensitive terms found in the document (up to 10)"],
        ],
        col_widths=[5, 12],
    )

    h(doc, "09. Exporting Results", 1)
    h(doc, "9.1 Export to CSV", 2)
    p(doc, '1. Run a complete scan.')
    p(doc, '2. Click the "Export CSV" button in the top right corner.')
    p(doc, '3. Choose location and filename in the save dialog.')
    p(doc, '4. Suggested name: osint_findings_YYYYMMDD_HHMM.csv')
    p(doc, 'Exported CSV columns:', bold=True)
    table(doc,
        ["CSV Column", "Description"],
        [
            ["risk_score", "Numeric score (1-100)"],
            ["risk_level", "Level: critical, high, medium, low"],
            ["country", "Country code (BR, US, FR, etc.)"],
            ["language", "Language code (pt, en, es, fr)"],
            ["entity_matched", "Detected entities"],
            ["source_platform", "Source platform"],
            ["title", "Full title"],
            ["file_type", "File format"],
            ["author", "Metadata author"],
            ["publisher", "Document publisher/creator"],
            ["cpf_count", "Number of CPFs"],
            ["cnpj_count", "Number of CNPJs"],
            ["financial_count", "Number of financial terms"],
            ["category", "Functional category"],
            ["url", "Full access URL"],
        ],
        col_widths=[4, 13],
    )

    h(doc, "10. Building the Executable (.exe)", 1)
    p(doc,
        'To distribute the application without requiring Python, '
        'use cx_Freeze to generate a standalone executable.'
    )
    h(doc, "10.1 Prerequisites", 2)
    code(doc, "pip install cx_Freeze")
    h(doc, "10.2 Build", 2)
    code(doc, "python setup_desktop.py build")
    p(doc,
        'The script automatically cleans build/ and dist/ folders before compiling. '
        'Output is generated in dist/OSINT_DLP_Desktop/.'
    )
    h(doc, "10.3 Generated Content", 2)
    table(doc,
        ["File/Folder", "Description"],
        [
            ["OSINT_DLP_Desktop.exe", "Main executable"],
            ["data/entities.json", "Group entity base (included automatically)"],
            ["lib/", "Bundled Python libraries"],
            ["python3XX.dll", "Embedded Python runtime"],
        ],
        col_widths=[5, 12],
    )
    h(doc, "10.4 Included vs. Excluded Packages", 2)
    table(doc,
        ["Included", "Excluded (Server only)"],
        [
            ["asyncio, json, csv, logging, threading", "fastapi"],
            ["tkinter, tkinter.ttk", "uvicorn, starlette"],
            ["httpx, fitz (PyMuPDF)", "jinja2"],
            ["docx (python-docx), openpyxl", "apscheduler"],
            ["ddgs, pydantic, pydantic_settings", "multipart"],
            ["sqlalchemy, aiosqlite", "test, unittest, pytest"],
        ],
        col_widths=[8.5, 8.5],
    )

    h(doc, "11. Key Files", 1)
    table(doc,
        ["File", "Lines", "Function"],
        [
            ["desktop/app.py", "436", "All GUI logic: interface, scan, export, details"],
            ["desktop/run_desktop.py", "10", "Entry point: resolves paths and calls desktop.app.main()"],
            ["setup_desktop.py", "47", "cx_Freeze build configuration for .exe"],
            ["data/entities.json", "~320", "Entity base (shared with Server)"],
            ["config/settings.py", "~40", "Environment settings (loads .env)"],
            ["config/intelligence_matrix.py", "~80", "Loads entities.json and builds intelligence matrix"],
            ["requirements.txt", "~15", "Python dependencies (shared)"],
        ],
        col_widths=[5, 2, 10],
    )

    h(doc, "12. Troubleshooting", 1)
    table(doc,
        ["Problem", "Likely Cause", "Solution"],
        [
            ["Window doesn't open", "Python without Tkinter", "Windows: reinstall Python with tkinter option. Linux: sudo apt install python3-tk"],
            ["Few results", "DuckDuckGo rate-limiting", "Wait 5 minutes and try again. Consider using PROXY_URL in .env"],
            ["Encoding error", "Console with limited charset", "Use terminal with UTF-8 support (Windows Terminal recommended)"],
            ["Scan takes long", "Many URLs to inspect", "Normal \u2014 each URL takes 1-3 seconds. ~80 dorks + delays = ~5-10 min total"],
            ["CSV button grayed", "No results in table", "Run a complete scan before exporting"],
            [".exe won't open", "Antivirus blocking", "Add exception in antivirus for the OSINT_DLP_Desktop directory"],
            ["ModuleNotFoundError", "Dependency not installed", "Run pip install -r requirements.txt in virtual environment"],
            ["entities.json not found", "Wrong relative path", "Always run from project root (cd Doc-Tracker)"],
        ],
        col_widths=[3, 4, 10],
    )

    h(doc, "13. Security and Privacy", 1)
    table(doc,
        ["Measure", "Description"],
        [
            ["Zero Disk", "All documents downloaded into BytesIO (memory) and never written to filesystem"],
            ["No Persistence", "Results exist only in application memory, discarded on close"],
            ["PII Masked", "CPFs, CNPJs, and emails masked in snippets (e.g., ***.456.789-**)"],
            ["Excluded Domains", "Official group sites automatically excluded from scanning"],
            ["Separate Thread", "Scan runs in separate thread to keep interface responsive"],
            ["No Local Network", "Application opens no network ports \u2014 zero attack surface"],
            ["Anti-block Delays", "Randomized intervals (2-4s) between queries to avoid rate-limiting"],
        ],
        col_widths=[4, 13],
    )

    end_page(doc, "Confidential Document\nRoullier Group Security Team\nVersion 2.1 \u2014 2026")
    return doc


def docx_to_pdf(docx_path, pdf_path):
    doc_in = Document(str(docx_path))
    pdf = fitz.open()
    pw, ph = fitz.paper_size("A4")
    m = 56
    uw = pw - 2 * m

    lines = []
    for para in doc_in.paragraphs:
        t = para.text.strip()
        if not t:
            lines.append(("blank", "", 8))
            continue
        sn = para.style.name if para.style else ""
        if any(r.bold and r.font.size and r.font.size >= Pt(28) for r in para.runs if r.text.strip()):
            lines.append(("cover_title", t, 40))
        elif any(r.font.size and r.font.size >= Pt(16) for r in para.runs if r.text.strip()):
            lines.append(("cover_sub", t, 24))
        elif "Heading 1" in sn:
            lines.append(("h1", t, 26))
        elif "Heading 2" in sn:
            lines.append(("h2", t, 20))
        elif "Heading 3" in sn:
            lines.append(("h3", t, 18))
        elif "No Spacing" in sn or any(r.font.name == "Consolas" for r in para.runs):
            for cl in t.split("\n"):
                lines.append(("code", cl, 13))
        elif "Bullet" in sn or "List" in sn:
            lines.append(("bullet", t, 15))
        elif t.startswith("____"):
            lines.append(("line", t, 10))
        else:
            lines.append(("body", t, 15))

    for tbl in doc_in.tables:
        lines.append(("blank", "", 4))
        for ri, row in enumerate(tbl.rows):
            rt = " | ".join(c.text.strip().replace("\n", " ") for c in row.cells)
            lines.append(("table_header" if ri == 0 else "table_row", rt, 14 if ri == 0 else 13))
        lines.append(("blank", "", 4))

    page = pdf.new_page(width=pw, height=ph)
    y = m
    on_cover = True

    for kind, text, height in lines:
        if kind == "cover_title" and on_cover:
            y = ph * 0.30
            page.insert_text((m, y), text, fontsize=28, fontname="hebo", color=(0.12, 0.25, 0.69))
            y += 10
            page.draw_rect(fitz.Rect(m, y, pw - m, y + 2), color=(0.02, 0.84, 0.63), fill=(0.02, 0.84, 0.63))
            y += 20
            continue
        elif kind == "cover_sub" and on_cover:
            page.insert_text((m, y), text, fontsize=13, fontname="helv", color=(0.39, 0.45, 0.55))
            y += height
            continue
        elif kind == "blank" and on_cover and y > ph * 0.5:
            page = pdf.new_page(width=pw, height=ph)
            y = m
            on_cover = False
            continue

        if y + height > ph - m:
            page = pdf.new_page(width=pw, height=ph)
            y = m

        if kind == "blank":
            y += height
        elif kind == "line":
            page.draw_rect(fitz.Rect(m, y - 2, pw - m, y), color=(0.02, 0.84, 0.63), fill=(0.02, 0.84, 0.63))
            y += 10
        elif kind == "h1":
            y += 10
            page.insert_text((m, y), text, fontsize=16, fontname="hebo", color=(0.12, 0.25, 0.69))
            y += 6
            page.draw_rect(fitz.Rect(m, y, m + 80, y + 2), color=(0.02, 0.84, 0.63), fill=(0.02, 0.84, 0.63))
            y += height
        elif kind == "h2":
            y += 6
            page.insert_text((m, y), text, fontsize=13, fontname="hebo", color=(0.12, 0.25, 0.69))
            y += height
        elif kind == "h3":
            y += 4
            page.insert_text((m, y), text, fontsize=11, fontname="hebo", color=(0.23, 0.35, 0.55))
            y += height
        elif kind == "code":
            bg = fitz.Rect(m + 10, y - 10, pw - m, y + 4)
            page.draw_rect(bg, color=None, fill=(0.94, 0.96, 0.98))
            page.insert_text((m + 14, y), text, fontsize=8.5, fontname="cour", color=(0.12, 0.16, 0.23))
            y += 13
        elif kind == "bullet":
            page.insert_text((m + 10, y), "\u2022", fontsize=10, fontname="helv", color=(0.02, 0.84, 0.63))
            words = text.split()
            cx = m + 24
            for word in words:
                tw = fitz.get_text_length(word + " ", fontsize=10, fontname="helv")
                if cx + tw > pw - m:
                    y += 13
                    cx = m + 24
                    if y > ph - m:
                        page = pdf.new_page(width=pw, height=ph)
                        y = m
                page.insert_text((cx, y), word + " ", fontsize=10, fontname="helv", color=(0.12, 0.16, 0.22))
                cx += tw
            y += height
        elif kind == "table_header":
            bg = fitz.Rect(m, y - 10, pw - m, y + 4)
            page.draw_rect(bg, color=None, fill=(0.06, 0.09, 0.16))
            page.insert_text((m + 4, y), text[:120], fontsize=8, fontname="hebo", color=(1, 1, 1))
            y += height
        elif kind == "table_row":
            if y + 14 > ph - m:
                page = pdf.new_page(width=pw, height=ph)
                y = m
            page.insert_text((m + 4, y), text[:140], fontsize=8, fontname="helv", color=(0.12, 0.16, 0.22))
            y += height
        else:
            words = text.split()
            cx = m
            for word in words:
                tw = fitz.get_text_length(word + " ", fontsize=10, fontname="helv")
                if cx + tw > pw - m:
                    y += 14
                    cx = m
                    if y > ph - m:
                        page = pdf.new_page(width=pw, height=ph)
                        y = m
                page.insert_text((cx, y), word + " ", fontsize=10, fontname="helv", color=(0.12, 0.16, 0.22))
                cx += tw
            y += height

    for i in range(len(pdf)):
        pg = pdf[i]
        pg.insert_text((pw - m - 30, ph - 20), f"{i+1}/{len(pdf)}", fontsize=8, fontname="helv", color=(0.58, 0.64, 0.72))

    pdf.save(str(pdf_path))
    pdf.close()


def main():
    print("[*] Desktop docs PT-BR (DOCX)...")
    doc_pt = build_pt()
    pt_docx = DOCS_DIR / "OSINT_DLP_Desktop_PT-BR.docx"
    doc_pt.save(str(pt_docx))
    print(f"    -> {pt_docx.name}")

    print("[*] Desktop docs EN (DOCX)...")
    doc_en = build_en()
    en_docx = DOCS_DIR / "OSINT_DLP_Desktop_EN.docx"
    doc_en.save(str(en_docx))
    print(f"    -> {en_docx.name}")

    print("[*] PDF PT-BR...")
    pt_pdf = DOCS_DIR / "OSINT_DLP_Desktop_PT-BR.pdf"
    docx_to_pdf(pt_docx, pt_pdf)
    print(f"    -> {pt_pdf.name}")

    print("[*] PDF EN...")
    en_pdf = DOCS_DIR / "OSINT_DLP_Desktop_EN.pdf"
    docx_to_pdf(en_docx, en_pdf)
    print(f"    -> {en_pdf.name}")

    print("[+] 4 documentos desktop gerados com sucesso!")


if __name__ == "__main__":
    main()
